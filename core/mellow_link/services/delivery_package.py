from __future__ import annotations

import hashlib
import json
import os
import re
import secrets
import uuid
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from enum import Enum
from pathlib import Path
from typing import Callable

from docx import Document
from pydantic import BaseModel, ConfigDict, Field, field_validator
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from sqlalchemy.orm import Session

from mellow_link.infra.database import (
    AnalysisContext,
    DeliveryAuditEvent,
    DeliveryChecklist,
    DeliveryChecklistItem,
    DeliveryChecklistTemplate,
    DeliveryChecklistTemplateItem,
    DeliveryCommandResult,
    DeliveryDownloadReference,
    DeliveryPackage,
    DeliveryPackageAssembly,
    ModernizationProject,
    PilotStateRecord,
    ProjectAsset,
    User,
    UserRole,
)
from mellow_link.services.project_results.archive import (
    build_project_result_archive_paths,
)

MIB = 1024 * 1024


def _bounded_limit(name: str, default: int, hard_max: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return default
    try:
        value = int(raw)
    except ValueError:
        return default
    return value if 0 < value <= hard_max else default


MAX_ARTIFACT_BYTES = _bounded_limit(
    "MELLOW_DELIVERY_MAX_ARTIFACT_BYTES", 25 * MIB, 25 * MIB
)
MAX_PACKAGE_INPUT_BYTES = _bounded_limit(
    "MELLOW_DELIVERY_MAX_INPUT_BYTES", 100 * MIB, 100 * MIB
)
MAX_PACKAGE_BYTES = _bounded_limit(
    "MELLOW_DELIVERY_MAX_PACKAGE_BYTES", 50 * MIB, 50 * MIB
)
MAX_ARTIFACT_COUNT = _bounded_limit("MELLOW_DELIVERY_MAX_ARTIFACTS", 20, 20)
MAX_DELIVERY_NOTE_BYTES = _bounded_limit(
    "MELLOW_DELIVERY_MAX_NOTE_BYTES", 64 * 1024, 64 * 1024
)
MAX_IDEMPOTENCY_KEY_LENGTH = 200
MAX_WAIVER_REASON_LENGTH = 1000
DOWNLOAD_REFERENCE_TTL = timedelta(minutes=15)
STALE_ASSEMBLY_AFTER = timedelta(minutes=10)
STAGING_RETENTION = timedelta(hours=24)
MAX_ASSEMBLY_ATTEMPTS = 3
MANIFEST_VERSION = "delivery-manifest-v1"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_REQUIRED_EXTERNAL_SECTIONS = tuple(f"{index}." for index in range(1, 10))
_FORBIDDEN_EXTERNAL_MARKERS = (
    "9. 분석 근거와 provenance",
    "run_id:",
    "safe_bundle_id",
    "original_path",
    "mapping_path",
)


class ChecklistItemStatus(str, Enum):
    PENDING = "pending"
    PRESENT = "present"
    MISSING = "missing"
    WAIVED = "waived"
    INVALID = "invalid"
    STALE = "stale"


class AssemblyStatus(str, Enum):
    PENDING = "pending"
    ASSEMBLING = "assembling"
    ASSEMBLED = "assembled"
    FAILED = "failed"
    SUPERSEDED = "superseded"


class ReadinessStatus(str, Enum):
    READY = "ready"
    NOT_READY = "not_ready"
    STALE = "stale"


class IdempotentRequest(BaseModel):
    idempotency_key: str = Field(min_length=1, max_length=MAX_IDEMPOTENCY_KEY_LENGTH)

    model_config = ConfigDict(extra="forbid")

    @field_validator("idempotency_key")
    @classmethod
    def validate_key(cls, value: str) -> str:
        if value != value.strip() or any(ord(char) < 32 for char in value):
            raise ValueError("idempotency_key contains invalid characters")
        return value


class VersionedChecklistRequest(IdempotentRequest):
    expected_checklist_version: int = Field(ge=0)


class WaiveChecklistItemRequest(VersionedChecklistRequest):
    reason: str = Field(min_length=1, max_length=MAX_WAIVER_REASON_LENGTH)

    @field_validator("reason")
    @classmethod
    def validate_reason(cls, value: str) -> str:
        if not value.strip() or any(ord(char) < 32 for char in value):
            raise ValueError("reason is invalid")
        return value


class RequestAssemblyRequest(IdempotentRequest):
    expected_pilot_version: int = Field(ge=0)
    expected_checklist_version: int = Field(ge=0)
    artifact_set_fingerprint: str = Field(pattern=r"^[0-9a-f]{64}$")
    manifest_version: str = Field(default=MANIFEST_VERSION)


class RetryAssemblyRequest(IdempotentRequest):
    expected_assembly_version: int = Field(ge=0)


class ChecklistItemView(BaseModel):
    item_key: str
    display_name: str
    requirement: str
    artifact_type: str
    status: ChecklistItemStatus
    artifact_ref: str | None = None
    artifact_fingerprint: str | None = None
    byte_size: int | None = None
    checksum: str | None = None
    content_type: str | None = None
    verified_at: datetime | None = None
    waived: bool
    version: int


class ChecklistView(BaseModel):
    checklist_id: str
    pilot_ref: str
    template_version: int
    version: int
    created_at: datetime
    updated_at: datetime
    items: list[ChecklistItemView]


class ReadinessIssue(BaseModel):
    item_key: str | None = None
    code: str


class ReadinessView(BaseModel):
    pilot_ref: str
    pilot_status: str
    pilot_version: int
    checklist_id: str
    checklist_version: int
    template_version: int
    readiness: ReadinessStatus
    blocking_items: list[ReadinessIssue]
    warnings: list[ReadinessIssue]
    artifact_set_fingerprint: str
    evaluated_at: datetime


class AssemblyView(BaseModel):
    assembly_id: str
    pilot_ref: str
    status: AssemblyStatus
    version: int
    attempt: int
    failure_code: str | None = None
    package_ref: str | None = None
    created_at: datetime
    updated_at: datetime
    completed_at: datetime | None = None


class PackageSummary(BaseModel):
    package_id: str
    package_ref: str
    status: str
    manifest_version: str
    byte_size: int
    checksum: str
    created_at: datetime


class PackagePage(BaseModel):
    items: list[PackageSummary]


class ManifestView(BaseModel):
    package_ref: str
    checksum: str
    byte_size: int
    manifest: dict[str, object]


class DownloadReferenceView(BaseModel):
    download_reference: str
    expires_at: datetime
    checksum: str
    byte_size: int


class DeliveryAuditView(BaseModel):
    event_id: str
    event_type: str
    actor_ref: str
    occurred_at: datetime
    result_version: int


class DeliveryAuditPage(BaseModel):
    items: list[DeliveryAuditView]


@dataclass(frozen=True)
class ResolvedDownload:
    path: Path
    filename: str
    content_type: str


@dataclass(frozen=True)
class _ResolvedArtifact:
    path: Path
    byte_size: int
    checksum: str
    fingerprint: str
    artifact_ref: str
    content_type: str


@dataclass(frozen=True)
class _TemplateItemSpec:
    item_key: str
    display_name: str
    description: str
    requirement: str
    artifact_type: str
    source: str
    waiver_allowed: bool = False


DEFAULT_TEMPLATE_KEY = "delivery-v1"
DEFAULT_TEMPLATE_VERSION = 1
DEFAULT_TEMPLATE_ITEMS = (
    _TemplateItemSpec(
        "external_report_docx",
        "External pilot report",
        "Reopenable external-safe DOCX report.",
        "required",
        "external_docx",
        "project_result_archive",
    ),
    _TemplateItemSpec(
        "one_page_summary",
        "One-page summary",
        "The fixed one-page summary section in the external report.",
        "required",
        "one_page_summary",
        "project_result_archive",
    ),
    _TemplateItemSpec(
        "provenance_summary",
        "External provenance summary",
        "The external-safe output criteria section.",
        "required",
        "external_provenance",
        "project_result_archive",
    ),
    _TemplateItemSpec(
        "delivery_note",
        "Delivery note",
        "Optional bounded operator note.",
        "optional",
        "delivery_note",
        "delivery_metadata",
    ),
)


class DeliveryError(Exception):
    code = "delivery_error"


class DeliveryNotFoundError(DeliveryError):
    code = "delivery_not_found"


class DeliveryAccessDeniedError(DeliveryError):
    code = "delivery_access_denied"


class DeliveryValidationError(DeliveryError):
    code = "delivery_validation_error"


class DeliveryVersionConflictError(DeliveryError):
    code = "delivery_version_conflict"

    def __init__(self, current_version: int):
        super().__init__("The resource version does not match expected_version")
        self.current_version = current_version


class DeliveryStateConflictError(DeliveryError):
    code = "delivery_state_conflict"


class ReadinessBlockedError(DeliveryError):
    code = "readiness_blocked"


class ReadinessStaleError(DeliveryError):
    code = "readiness_stale"


class DeliveryIdempotencyConflictError(DeliveryError):
    code = "idempotency_key_reused"


class AssemblyAlreadyExistsError(DeliveryError):
    code = "assembly_already_exists"


class InvalidArtifactError(DeliveryError):
    code = "invalid_artifact"


class PackageIntegrityError(DeliveryError):
    code = "package_integrity_failed"


class DeliveryStorageError(DeliveryError):
    code = "delivery_storage_error"


class DeliveryPackageService:
    def __init__(
        self,
        db: Session,
        *,
        archive_root: Path | None = None,
        now_provider: Callable[[], datetime] | None = None,
    ):
        self.db = db
        self.archive_root = archive_root or (
            Path(__file__).resolve().parents[3]
            / "data"
            / "outputs"
            / "final"
            / "project_results"
        )
        self.now_provider = now_provider or (lambda: datetime.now(timezone.utc))

    def create_checklist(
        self, actor: User, pilot_id: str, payload: IdempotentRequest
    ) -> ChecklistView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        if pilot.status == "delivered":
            raise DeliveryStateConflictError("Delivered Pilot is read-only")
        request_hash = _request_hash("create_checklist", {"pilot_id": pilot_id})
        replay = self._load_replay(
            actor.id,
            payload.idempotency_key,
            "create_checklist",
            request_hash,
            ChecklistView,
        )
        if replay is not None:
            return replay
        template = self._ensure_template()
        existing = self._find_checklist(pilot_id, template.template_id)
        if existing is not None:
            return self._checklist_view(existing)

        now = self._now()
        checklist = DeliveryChecklist(
            checklist_id=uuid.uuid4().hex,
            pilot_id=pilot.pilot_id,
            project_id=pilot.project_id,
            run_id=pilot.run_id,
            template_id=template.template_id,
            template_version=template.template_version,
            version=0,
            created_at=now,
            updated_at=now,
            created_by_id=actor.id,
        )
        self.db.add(checklist)
        try:
            self.db.flush()
            template_items = (
                self.db.query(DeliveryChecklistTemplateItem)
                .filter(
                    DeliveryChecklistTemplateItem.template_id == template.template_id
                )
                .order_by(DeliveryChecklistTemplateItem.sort_order.asc())
                .all()
            )
            for item in template_items:
                self.db.add(
                    DeliveryChecklistItem(
                        checklist_item_id=uuid.uuid4().hex,
                        checklist_id=checklist.checklist_id,
                        item_key=item.item_key,
                        display_name=item.display_name,
                        description=item.description,
                        requirement=item.requirement,
                        artifact_type=item.artifact_type,
                        source=item.source,
                        waiver_allowed=item.waiver_allowed,
                        sort_order=item.sort_order,
                        status=ChecklistItemStatus.PENDING.value,
                        version=0,
                        created_at=now,
                        updated_at=now,
                    )
                )
            self.db.flush()
            self._add_audit(
                pilot,
                actor,
                "checklist_created",
                now,
                checklist_id=checklist.checklist_id,
                idempotency_key=payload.idempotency_key,
                result_version=0,
            )
            response = self._checklist_view(checklist)
            self._store_result(
                actor.id,
                payload.idempotency_key,
                "create_checklist",
                request_hash,
                "checklist",
                checklist.checklist_id,
                0,
                response,
                now,
            )
            self.db.commit()
            return response
        except IntegrityError as exc:
            self.db.rollback()
            replay = self._load_replay(
                actor.id,
                payload.idempotency_key,
                "create_checklist",
                request_hash,
                ChecklistView,
            )
            if replay is not None:
                return replay
            template = self._current_template()
            existing = self._find_checklist(pilot_id, template.template_id)
            if existing is not None:
                return self._checklist_view(existing)
            raise DeliveryStorageError("Checklist could not be created") from exc
        except SQLAlchemyError as exc:
            self.db.rollback()
            raise DeliveryStorageError("Checklist could not be created") from exc

    def get_checklist(self, actor: User, pilot_id: str) -> ChecklistView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_project_access(actor, project)
        checklist = self._get_checklist(pilot.pilot_id)
        return self._checklist_view(checklist, derive_stale=True)

    def get_item(self, actor: User, pilot_id: str, item_key: str) -> ChecklistItemView:
        checklist = self.get_checklist(actor, pilot_id)
        for item in checklist.items:
            if item.item_key == item_key:
                return item
        raise DeliveryNotFoundError("Checklist item was not found")

    def verify_item(
        self,
        actor: User,
        pilot_id: str,
        item_key: str,
        payload: VersionedChecklistRequest,
    ) -> ChecklistView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        self._require_mutable(pilot)
        checklist = self._get_checklist(pilot_id)
        item = self._get_item_record(checklist.checklist_id, item_key)
        operation = "verify_checklist_item"
        request_hash = _request_hash(
            operation,
            {
                "pilot_id": pilot_id,
                "item_key": item_key,
                "expected_checklist_version": payload.expected_checklist_version,
            },
        )
        replay = self._load_replay(
            actor.id,
            payload.idempotency_key,
            operation,
            request_hash,
            ChecklistView,
        )
        if replay is not None:
            return replay
        self._require_checklist_version(checklist, payload.expected_checklist_version)
        self._require_not_frozen(checklist.checklist_id)
        artifact, status = self._resolve_item(pilot, item)
        now = self._now()
        new_version = checklist.version + 1
        updated = (
            self.db.query(DeliveryChecklist)
            .filter(
                DeliveryChecklist.checklist_id == checklist.checklist_id,
                DeliveryChecklist.version == payload.expected_checklist_version,
            )
            .update(
                {"version": new_version, "updated_at": now},
                synchronize_session=False,
            )
        )
        if updated != 1:
            self.db.rollback()
            current = self._get_checklist(pilot_id)
            raise DeliveryVersionConflictError(current.version)
        item.status = status.value
        item.artifact_ref = artifact.artifact_ref if artifact else None
        item.artifact_fingerprint = artifact.fingerprint if artifact else None
        item.verified_by_id = actor.id
        item.verified_at = now
        item.waived_by_id = None
        item.waived_at = None
        item.waiver_reason = None
        item.version += 1
        item.updated_at = now
        try:
            self.db.flush()
            self._add_audit(
                pilot,
                actor,
                "checklist_item_verified",
                now,
                checklist_id=checklist.checklist_id,
                idempotency_key=payload.idempotency_key,
                result_version=new_version,
            )
            self.db.flush()
            self.db.expire_all()
            response = self._checklist_view(self._get_checklist(pilot_id))
            self._store_result(
                actor.id,
                payload.idempotency_key,
                operation,
                request_hash,
                "checklist",
                checklist.checklist_id,
                new_version,
                response,
                now,
            )
            self.db.commit()
            return response
        except SQLAlchemyError as exc:
            self.db.rollback()
            raise DeliveryStorageError(
                "Checklist verification could not be stored"
            ) from exc

    def waive_item(
        self,
        actor: User,
        pilot_id: str,
        item_key: str,
        payload: WaiveChecklistItemRequest,
    ) -> ChecklistView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        self._require_mutable(pilot)
        checklist = self._get_checklist(pilot_id)
        item = self._get_item_record(checklist.checklist_id, item_key)
        if not item.waiver_allowed:
            raise DeliveryStateConflictError("Checklist item cannot be waived")
        operation = "waive_checklist_item"
        request_hash = _request_hash(
            operation,
            {
                "pilot_id": pilot_id,
                "item_key": item_key,
                "expected_checklist_version": payload.expected_checklist_version,
                "reason": payload.reason,
            },
        )
        replay = self._load_replay(
            actor.id,
            payload.idempotency_key,
            operation,
            request_hash,
            ChecklistView,
        )
        if replay is not None:
            return replay
        self._require_checklist_version(checklist, payload.expected_checklist_version)
        self._require_not_frozen(checklist.checklist_id)
        now = self._now()
        new_version = checklist.version + 1
        updated = (
            self.db.query(DeliveryChecklist)
            .filter(
                DeliveryChecklist.checklist_id == checklist.checklist_id,
                DeliveryChecklist.version == payload.expected_checklist_version,
            )
            .update(
                {"version": new_version, "updated_at": now},
                synchronize_session=False,
            )
        )
        if updated != 1:
            self.db.rollback()
            raise DeliveryVersionConflictError(self._get_checklist(pilot_id).version)
        item.status = ChecklistItemStatus.WAIVED.value
        item.waived_by_id = actor.id
        item.waived_at = now
        item.waiver_reason = payload.reason
        item.version += 1
        item.updated_at = now
        try:
            self.db.flush()
            self._add_audit(
                pilot,
                actor,
                "checklist_item_waived",
                now,
                checklist_id=checklist.checklist_id,
                idempotency_key=payload.idempotency_key,
                result_version=new_version,
            )
            self.db.flush()
            self.db.expire_all()
            response = self._checklist_view(self._get_checklist(pilot_id))
            self._store_result(
                actor.id,
                payload.idempotency_key,
                operation,
                request_hash,
                "checklist",
                checklist.checklist_id,
                new_version,
                response,
                now,
            )
            self.db.commit()
            return response
        except SQLAlchemyError as exc:
            self.db.rollback()
            raise DeliveryStorageError("Checklist waiver could not be stored") from exc

    def get_readiness(self, actor: User, pilot_id: str) -> ReadinessView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_project_access(actor, project)
        checklist = self._get_checklist(pilot_id)
        return self._readiness(pilot, checklist)

    def request_assembly(
        self, actor: User, pilot_id: str, payload: RequestAssemblyRequest
    ) -> AssemblyView:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        operation = "request_package_assembly"
        request_hash = _request_hash(
            operation,
            {
                "pilot_id": pilot_id,
                "expected_pilot_version": payload.expected_pilot_version,
                "expected_checklist_version": payload.expected_checklist_version,
                "artifact_set_fingerprint": payload.artifact_set_fingerprint,
                "manifest_version": payload.manifest_version,
            },
        )
        replay = self._load_replay(
            actor.id,
            payload.idempotency_key,
            operation,
            request_hash,
            AssemblyView,
        )
        if replay is not None:
            if replay.status == AssemblyStatus.PENDING:
                assembly = self._get_assembly(replay.assembly_id)
                return self._run_assembly(actor, assembly, payload.idempotency_key)
            return replay
        self.recover_interrupted_assemblies(actor)
        pilot = self._get_pilot(pilot_id)
        if pilot.status != "approved":
            raise DeliveryStateConflictError("Pilot must be approved")
        if pilot.version != payload.expected_pilot_version:
            raise DeliveryVersionConflictError(pilot.version)
        if payload.manifest_version != MANIFEST_VERSION:
            raise DeliveryValidationError("manifest_version is unsupported")
        checklist = self._get_checklist(pilot_id)
        self._require_checklist_version(checklist, payload.expected_checklist_version)
        readiness = self._readiness(pilot, checklist)
        if readiness.readiness == ReadinessStatus.STALE:
            raise ReadinessStaleError("Readiness snapshot is stale")
        if readiness.readiness != ReadinessStatus.READY:
            raise ReadinessBlockedError("Delivery readiness is blocked")
        if readiness.artifact_set_fingerprint != payload.artifact_set_fingerprint:
            raise ReadinessStaleError("Artifact fingerprint changed")
        request_payload = {
            "pilot_id": pilot_id,
            "source_pilot_version": pilot.version,
            "checklist_id": checklist.checklist_id,
            "checklist_version": checklist.version,
            "template_version": checklist.template_version,
            "artifact_set_fingerprint": readiness.artifact_set_fingerprint,
            "manifest_version": payload.manifest_version,
            "output_profile": "external-v1",
        }
        fingerprint = hashlib.sha256(
            json.dumps(request_payload, sort_keys=True, separators=(",", ":")).encode()
        ).hexdigest()
        existing = (
            self.db.query(DeliveryPackageAssembly)
            .filter(DeliveryPackageAssembly.request_fingerprint == fingerprint)
            .first()
        )
        if existing is not None:
            return self._assembly_view(existing)
        now = self._now()
        assembly = DeliveryPackageAssembly(
            assembly_id=uuid.uuid4().hex,
            pilot_id=pilot.pilot_id,
            checklist_id=checklist.checklist_id,
            project_id=pilot.project_id,
            run_id=pilot.run_id,
            status=AssemblyStatus.PENDING.value,
            version=0,
            attempt=1,
            request_fingerprint=fingerprint,
            source_pilot_version=pilot.version,
            checklist_version=checklist.version,
            template_version=checklist.template_version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            manifest_version=payload.manifest_version,
            created_by_id=actor.id,
            created_at=now,
            updated_at=now,
        )
        try:
            self.db.add(assembly)
            self.db.flush()
            self._add_audit(
                pilot,
                actor,
                "package_assembly_requested",
                now,
                checklist_id=checklist.checklist_id,
                assembly_id=assembly.assembly_id,
                idempotency_key=payload.idempotency_key,
                result_version=0,
            )
            pending = self._assembly_view(assembly)
            self._store_result(
                actor.id,
                payload.idempotency_key,
                operation,
                request_hash,
                "assembly",
                assembly.assembly_id,
                0,
                pending,
                now,
            )
            self.db.commit()
        except IntegrityError as exc:
            self.db.rollback()
            replay = self._load_replay(
                actor.id,
                payload.idempotency_key,
                operation,
                request_hash,
                AssemblyView,
            )
            if replay is not None:
                return replay
            existing = (
                self.db.query(DeliveryPackageAssembly)
                .filter(DeliveryPackageAssembly.request_fingerprint == fingerprint)
                .first()
            )
            if existing is not None:
                return self._assembly_view(existing)
            raise DeliveryStorageError("Assembly request could not be stored") from exc
        return self._run_assembly(actor, assembly, payload.idempotency_key)

    def get_assembly(self, actor: User, assembly_id: str) -> AssemblyView:
        assembly = self._get_assembly(assembly_id)
        project = self._get_project(assembly.project_id)
        self._require_project_access(actor, project)
        return self._assembly_view(assembly)

    def retry_assembly(
        self, actor: User, assembly_id: str, payload: RetryAssemblyRequest
    ) -> AssemblyView:
        assembly = self._get_assembly(assembly_id)
        pilot, project = self._pilot_and_project(assembly.pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        self.recover_interrupted_assemblies(actor)
        assembly = self._get_assembly(assembly_id)
        if assembly.version != payload.expected_assembly_version:
            raise DeliveryVersionConflictError(assembly.version)
        if assembly.status != AssemblyStatus.FAILED.value:
            raise DeliveryStateConflictError("Only failed assembly can be retried")
        if assembly.attempt >= MAX_ASSEMBLY_ATTEMPTS:
            raise DeliveryStateConflictError("Assembly retry limit reached")
        if assembly.failure_code in {
            "invalid_artifact",
            "readiness_stale",
            "package_size_exceeded",
        }:
            raise DeliveryStateConflictError("Assembly requires a new request")
        if pilot.status != "approved":
            raise DeliveryStateConflictError("Pilot must remain approved")
        checklist = self._get_checklist(pilot.pilot_id)
        readiness = self._readiness(pilot, checklist)
        if (
            readiness.readiness != ReadinessStatus.READY
            or readiness.artifact_set_fingerprint != assembly.artifact_set_fingerprint
        ):
            raise ReadinessStaleError("Assembly source is no longer current")
        operation = "retry_package_assembly"
        request_hash = _request_hash(
            operation,
            {
                "assembly_id": assembly_id,
                "expected_assembly_version": payload.expected_assembly_version,
            },
        )
        replay = self._load_replay(
            actor.id,
            payload.idempotency_key,
            operation,
            request_hash,
            AssemblyView,
        )
        if replay is not None:
            return replay
        now = self._now()
        updated = (
            self.db.query(DeliveryPackageAssembly)
            .filter(
                DeliveryPackageAssembly.assembly_id == assembly_id,
                DeliveryPackageAssembly.version == payload.expected_assembly_version,
                DeliveryPackageAssembly.status == AssemblyStatus.FAILED.value,
            )
            .update(
                {
                    "status": AssemblyStatus.PENDING.value,
                    "version": assembly.version + 1,
                    "attempt": assembly.attempt + 1,
                    "updated_at": now,
                    "failure_code": None,
                    "failed_at": None,
                },
                synchronize_session=False,
            )
        )
        if updated != 1:
            self.db.rollback()
            raise DeliveryVersionConflictError(self._get_assembly(assembly_id).version)
        self.db.expire_all()
        current = self._get_assembly(assembly_id)
        pending = self._assembly_view(current)
        try:
            self._add_audit(
                pilot,
                actor,
                "package_assembly_retried",
                now,
                checklist_id=assembly.checklist_id,
                assembly_id=assembly_id,
                idempotency_key=payload.idempotency_key,
                result_version=current.version,
            )
            self._store_result(
                actor.id,
                payload.idempotency_key,
                operation,
                request_hash,
                "assembly",
                assembly_id,
                current.version,
                pending,
                now,
            )
            self.db.commit()
        except SQLAlchemyError as exc:
            self.db.rollback()
            raise DeliveryStorageError("Assembly retry could not be stored") from exc
        return self._run_assembly(actor, current, payload.idempotency_key)

    def get_manifest(self, actor: User, package_id: str) -> ManifestView:
        package = self._get_package(package_id)
        self._require_project_access(actor, self._get_project(package.project_id))
        self._verify_package_file(package)
        return ManifestView(
            package_ref=_safe_ref("package", package.package_id),
            checksum=package.checksum,
            byte_size=package.byte_size,
            manifest=json.loads(package.manifest_json),
        )

    def list_packages(self, actor: User, pilot_id: str) -> PackagePage:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_project_access(actor, project)
        packages = (
            self.db.query(DeliveryPackage)
            .filter(DeliveryPackage.pilot_id == pilot.pilot_id)
            .order_by(
                DeliveryPackage.created_at.desc(), DeliveryPackage.package_id.desc()
            )
            .limit(100)
            .all()
        )
        return PackagePage(items=[self._package_summary(item) for item in packages])

    def create_download_reference(
        self, actor: User, package_id: str
    ) -> DownloadReferenceView:
        package = self._get_package(package_id)
        self._require_operator(actor)
        self._require_project_access(actor, self._get_project(package.project_id))
        self._verify_package_file(package)
        token = secrets.token_urlsafe(32)
        now = self._now()
        reference = DeliveryDownloadReference(
            reference_id=uuid.uuid4().hex,
            package_id=package.package_id,
            actor_id=actor.id,
            token_digest=_sha256_bytes(token.encode()),
            created_at=now,
            expires_at=now + DOWNLOAD_REFERENCE_TTL,
        )
        try:
            self.db.add(reference)
            pilot = self._get_pilot(package.pilot_id)
            self._add_audit(
                pilot,
                actor,
                "package_download_reference_issued",
                now,
                package_id=package.package_id,
                result_version=0,
            )
            self.db.commit()
        except SQLAlchemyError as exc:
            self.db.rollback()
            raise DeliveryStorageError(
                "Download reference could not be stored"
            ) from exc
        return DownloadReferenceView(
            download_reference=token,
            expires_at=now + DOWNLOAD_REFERENCE_TTL,
            checksum=package.checksum,
            byte_size=package.byte_size,
        )

    def redeem_download_reference(self, actor: User, token: str) -> ResolvedDownload:
        if not token or len(token) > 200 or any(ord(char) < 33 for char in token):
            raise DeliveryNotFoundError("Download reference was not found")
        digest = _sha256_bytes(token.encode())
        reference = (
            self.db.query(DeliveryDownloadReference)
            .filter(DeliveryDownloadReference.token_digest == digest)
            .first()
        )
        if reference is None:
            raise DeliveryNotFoundError("Download reference was not found")
        self._require_operator(actor)
        if reference.actor_id != actor.id:
            raise DeliveryAccessDeniedError("Download reference is not authorized")
        now = self._now()
        if reference.consumed_at is not None or _as_utc(reference.expires_at) <= now:
            raise DeliveryStateConflictError("Download reference is no longer valid")
        package = self._get_package(reference.package_id)
        self._require_project_access(actor, self._get_project(package.project_id))
        path = self._verify_package_file(package)
        updated = (
            self.db.query(DeliveryDownloadReference)
            .filter(
                DeliveryDownloadReference.reference_id == reference.reference_id,
                DeliveryDownloadReference.consumed_at.is_(None),
            )
            .update({"consumed_at": now}, synchronize_session=False)
        )
        if updated != 1:
            self.db.rollback()
            raise DeliveryStateConflictError("Download reference is no longer valid")
        self.db.commit()
        return ResolvedDownload(
            path=path,
            filename="pilot-delivery-package.zip",
            content_type="application/zip",
        )

    def get_audit_history(self, actor: User, pilot_id: str) -> DeliveryAuditPage:
        pilot, project = self._pilot_and_project(pilot_id)
        self._require_operator(actor)
        self._require_project_access(actor, project)
        events = (
            self.db.query(DeliveryAuditEvent)
            .filter(DeliveryAuditEvent.pilot_id == pilot.pilot_id)
            .order_by(
                DeliveryAuditEvent.occurred_at.asc(), DeliveryAuditEvent.event_id.asc()
            )
            .all()
        )
        return DeliveryAuditPage(
            items=[
                DeliveryAuditView(
                    event_id=event.event_id,
                    event_type=event.event_type,
                    actor_ref=_safe_ref("actor", str(event.actor_id)),
                    occurred_at=_as_utc(event.occurred_at),
                    result_version=event.result_version,
                )
                for event in events
            ]
        )

    def recover_interrupted_assemblies(self, actor: User) -> int:
        self._require_operator(actor)
        now = self._now()
        cutoff = now - STALE_ASSEMBLY_AFTER
        records = (
            self.db.query(DeliveryPackageAssembly)
            .filter(
                DeliveryPackageAssembly.status == AssemblyStatus.ASSEMBLING.value,
                DeliveryPackageAssembly.updated_at < cutoff,
            )
            .all()
        )
        recovered = 0
        for assembly in records:
            pilot = self._get_pilot(assembly.pilot_id)
            assembly.status = AssemblyStatus.FAILED.value
            assembly.version += 1
            assembly.failed_at = now
            assembly.updated_at = now
            assembly.failure_code = "assembly_interrupted"
            self._add_audit(
                pilot,
                actor,
                "package_assembly_failed",
                now,
                checklist_id=assembly.checklist_id,
                assembly_id=assembly.assembly_id,
                result_version=assembly.version,
            )
            self.db.flush()
            response = self._assembly_view(assembly)
            for command in (
                self.db.query(DeliveryCommandResult)
                .filter(
                    DeliveryCommandResult.resource_type == "assembly",
                    DeliveryCommandResult.resource_id == assembly.assembly_id,
                )
                .all()
            ):
                command.result_version = assembly.version
                command.response_json = response.model_dump_json()
            recovered += 1
        self._cleanup_staging(now)
        if recovered:
            self.db.commit()
        return recovered

    def _run_assembly(
        self, actor: User, assembly: DeliveryPackageAssembly, idempotency_key: str
    ) -> AssemblyView:
        assembly = self._get_assembly(assembly.assembly_id)
        if assembly.status != AssemblyStatus.PENDING.value:
            return self._assembly_view(assembly)
        now = self._now()
        claim_version = assembly.version + 1
        claimed = (
            self.db.query(DeliveryPackageAssembly)
            .filter(
                DeliveryPackageAssembly.assembly_id == assembly.assembly_id,
                DeliveryPackageAssembly.status == AssemblyStatus.PENDING.value,
                DeliveryPackageAssembly.version == assembly.version,
            )
            .update(
                {
                    "status": AssemblyStatus.ASSEMBLING.value,
                    "version": claim_version,
                    "started_at": now,
                    "updated_at": now,
                },
                synchronize_session=False,
            )
        )
        if claimed != 1:
            self.db.rollback()
            return self._assembly_view(self._get_assembly(assembly.assembly_id))
        pilot = self._get_pilot(assembly.pilot_id)
        self._add_audit(
            pilot,
            actor,
            "package_assembly_started",
            now,
            checklist_id=assembly.checklist_id,
            assembly_id=assembly.assembly_id,
            idempotency_key=idempotency_key,
            result_version=claim_version,
        )
        self.db.commit()
        self.db.expire_all()
        assembly = self._get_assembly(assembly.assembly_id)
        try:
            package_id, manifest, final_path, checksum, byte_size = self._build_archive(
                actor, assembly
            )
            completed_at = self._now()
            prior_packages = (
                self.db.query(DeliveryPackage)
                .filter(
                    DeliveryPackage.pilot_id == assembly.pilot_id,
                    DeliveryPackage.status == "assembled",
                )
                .all()
            )
            for prior in prior_packages:
                prior.status = "superseded"
                prior.superseded_at = completed_at
                prior_assembly = self._get_assembly(prior.assembly_id)
                prior_assembly.status = AssemblyStatus.SUPERSEDED.value
                prior_assembly.version += 1
                prior_assembly.updated_at = completed_at
            package = DeliveryPackage(
                package_id=package_id,
                assembly_id=assembly.assembly_id,
                pilot_id=assembly.pilot_id,
                project_id=assembly.project_id,
                run_id=assembly.run_id,
                status="assembled",
                manifest_version=MANIFEST_VERSION,
                manifest_json=json.dumps(
                    manifest, ensure_ascii=True, sort_keys=True, separators=(",", ":")
                ),
                artifact_reference=package_id,
                byte_size=byte_size,
                checksum=checksum,
                created_by_id=actor.id,
                created_at=completed_at,
            )
            self.db.add(package)
            updated = (
                self.db.query(DeliveryPackageAssembly)
                .filter(
                    DeliveryPackageAssembly.assembly_id == assembly.assembly_id,
                    DeliveryPackageAssembly.status == AssemblyStatus.ASSEMBLING.value,
                    DeliveryPackageAssembly.version == claim_version,
                )
                .update(
                    {
                        "status": AssemblyStatus.ASSEMBLED.value,
                        "version": claim_version + 1,
                        "completed_at": completed_at,
                        "updated_at": completed_at,
                    },
                    synchronize_session=False,
                )
            )
            if updated != 1:
                raise DeliveryStorageError("Assembly completion conflicted")
            self._add_audit(
                pilot,
                actor,
                "package_assembled",
                completed_at,
                checklist_id=assembly.checklist_id,
                assembly_id=assembly.assembly_id,
                package_id=package_id,
                idempotency_key=idempotency_key,
                result_version=claim_version + 1,
            )
            self.db.flush()
            self.db.expire_all()
            response = self._assembly_view(self._get_assembly(assembly.assembly_id))
            self._replace_command_response(idempotency_key, actor.id, response)
            self.db.commit()
            return response
        except Exception as exc:
            self.db.rollback()
            failure_code = _safe_failure_code(exc)
            failed_at = self._now()
            assembly = self._get_assembly(assembly.assembly_id)
            if assembly.status == AssemblyStatus.ASSEMBLING.value:
                assembly.status = AssemblyStatus.FAILED.value
                assembly.version += 1
                assembly.failed_at = failed_at
                assembly.updated_at = failed_at
                assembly.failure_code = failure_code
                self._add_audit(
                    pilot,
                    actor,
                    "package_assembly_failed",
                    failed_at,
                    checklist_id=assembly.checklist_id,
                    assembly_id=assembly.assembly_id,
                    idempotency_key=idempotency_key,
                    result_version=assembly.version,
                )
                self.db.flush()
                response = self._assembly_view(assembly)
                self._replace_command_response(idempotency_key, actor.id, response)
                self.db.commit()
                return response
            raise DeliveryStorageError("Assembly failed") from exc

    def _build_archive(
        self, actor: User, assembly: DeliveryPackageAssembly
    ) -> tuple[str, dict[str, object], Path, str, int]:
        pilot = self._get_pilot(assembly.pilot_id)
        artifact = self._resolve_external_docx(pilot)
        if artifact.byte_size > MAX_ARTIFACT_BYTES:
            raise InvalidArtifactError("Artifact exceeds size policy")
        package_id = uuid.uuid4().hex
        created_at = self._now()
        artifact_entry = {
            "artifact_type": "external_docx",
            "entry_name": "report/pilot-report.docx",
            "byte_size": artifact.byte_size,
            "checksum_algorithm": "sha256",
            "checksum": artifact.checksum,
            "content_type": DOCX_CONTENT_TYPE,
            "source_fingerprint": artifact.fingerprint,
        }
        manifest: dict[str, object] = {
            "manifest_version": MANIFEST_VERSION,
            "package_id": _safe_ref("package", package_id),
            "pilot_ref": _safe_ref("pilot", assembly.pilot_id),
            "created_at": created_at.isoformat(),
            "created_by_ref": _safe_ref("actor", str(actor.id)),
            "source_pilot_version": assembly.source_pilot_version,
            "checklist_template_version": assembly.template_version,
            "checklist_version": assembly.checklist_version,
            "artifact_set_fingerprint": assembly.artifact_set_fingerprint,
            "artifacts": [artifact_entry],
            "logical_coverage": ["one_page_summary", "external_provenance"],
            "assembly_status": "assembled",
        }
        paths = build_project_result_archive_paths(
            archive_root=self.archive_root,
            project_id=assembly.project_id,
            run_id=assembly.run_id,
        )
        self._assert_safe_root(paths["dir"], self.archive_root)
        package_root = paths["dir"] / "delivery_packages"
        staging_root = package_root / ".staging"
        package_root.mkdir(parents=True, exist_ok=True)
        staging_root.mkdir(parents=True, exist_ok=True)
        self._assert_safe_root(package_root, paths["dir"])
        self._assert_safe_root(staging_root, package_root)
        staging_path = staging_root / f"{package_id}.partial"
        final_path = package_root / f"{package_id}.zip"
        if staging_path.exists() or final_path.exists():
            raise DeliveryStorageError("Package destination already exists")
        manifest_bytes = json.dumps(
            manifest,
            ensure_ascii=True,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
        if artifact.byte_size + len(manifest_bytes) > MAX_PACKAGE_INPUT_BYTES:
            raise InvalidArtifactError("Package input exceeds size policy")
        try:
            with zipfile.ZipFile(
                staging_path,
                mode="x",
                compression=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            ) as archive:
                _write_zip_entry(archive, "manifest.json", manifest_bytes)
                _write_zip_entry(
                    archive, "report/pilot-report.docx", artifact.path.read_bytes()
                )
            if staging_path.stat().st_size > MAX_PACKAGE_BYTES:
                raise InvalidArtifactError("Package exceeds size policy")
            self._verify_archive(staging_path, manifest)
            os.replace(staging_path, final_path)
            checksum = _sha256_path(final_path)
            return package_id, manifest, final_path, checksum, final_path.stat().st_size
        except Exception:
            staging_path.unlink(missing_ok=True)
            raise

    def _verify_archive(self, path: Path, manifest: dict[str, object]) -> None:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            if names != ["manifest.json", "report/pilot-report.docx"]:
                raise PackageIntegrityError("Package entries do not match allowlist")
            stored_manifest = json.loads(archive.read("manifest.json"))
            if stored_manifest != manifest:
                raise PackageIntegrityError("Manifest verification failed")
            report = archive.read("report/pilot-report.docx")
            expected = manifest["artifacts"][0]["checksum"]  # type: ignore[index]
            if _sha256_bytes(report) != expected:
                raise PackageIntegrityError("Artifact checksum verification failed")

    def _resolve_item(
        self, pilot: PilotStateRecord, item: DeliveryChecklistItem
    ) -> tuple[_ResolvedArtifact | None, ChecklistItemStatus]:
        if item.artifact_type == "delivery_note":
            return None, ChecklistItemStatus.MISSING
        try:
            artifact = self._resolve_external_docx(pilot)
            text = _docx_text(artifact.path)
            if (
                item.artifact_type == "one_page_summary"
                and "1. 1페이지 요약" not in text
            ):
                raise InvalidArtifactError("Required summary section is missing")
            if (
                item.artifact_type == "external_provenance"
                and "9. 산출물 기준" not in text
            ):
                raise InvalidArtifactError("External provenance section is missing")
            return artifact, ChecklistItemStatus.PRESENT
        except FileNotFoundError:
            return None, ChecklistItemStatus.MISSING
        except (
            InvalidArtifactError,
            PackageIntegrityError,
            zipfile.BadZipFile,
            ValueError,
        ):
            return None, ChecklistItemStatus.INVALID

    def _resolve_external_docx(self, pilot: PilotStateRecord) -> _ResolvedArtifact:
        paths = build_project_result_archive_paths(
            archive_root=self.archive_root,
            project_id=pilot.project_id,
            run_id=pilot.run_id,
        )
        self._assert_safe_root(paths["dir"], self.archive_root)
        path = paths["external_docx"]
        if not path.is_file():
            raise FileNotFoundError
        self._assert_safe_root(path, paths["dir"])
        current = path
        has_symlink = False
        while True:
            has_symlink = has_symlink or current.is_symlink()
            if current == paths["dir"] or current.parent == current:
                break
            current = current.parent
        if has_symlink:
            raise InvalidArtifactError("Symbolic links are not allowed")
        size = path.stat().st_size
        if size > MAX_ARTIFACT_BYTES:
            raise InvalidArtifactError("Artifact exceeds size policy")
        _validate_docx(path)
        text = _docx_text(path)
        if not all(
            any(line.startswith(section) for line in text.splitlines())
            for section in _REQUIRED_EXTERNAL_SECTIONS
        ):
            raise InvalidArtifactError("External report structure is incomplete")
        if any(marker in text for marker in _FORBIDDEN_EXTERNAL_MARKERS):
            raise InvalidArtifactError("External report boundary validation failed")
        source_names = [
            row[0]
            for row in self.db.query(ProjectAsset.original_filename)
            .filter(ProjectAsset.project_id == pilot.project_id)
            .all()
            if row[0]
        ]
        bundle_ids = [
            row[0]
            for row in self.db.query(AnalysisContext.safe_bundle_id)
            .filter(AnalysisContext.project_id == pilot.project_id)
            .all()
            if row[0]
        ]
        sensitive_identifiers = [
            value
            for value in [pilot.project_id, pilot.run_id, *source_names, *bundle_ids]
            if len(value) >= 8
        ]
        if any(value in text for value in sensitive_identifiers):
            raise InvalidArtifactError("External report contains internal identifiers")
        if re.search(r"(?:[A-Za-z]:[\\/]|\\\\|/(?:home|Users)/)", text):
            raise InvalidArtifactError("External report contains an internal path")
        checksum = _sha256_path(path)
        fingerprint = _sha256_bytes(f"docx-v1:{checksum}".encode())
        return _ResolvedArtifact(
            path=path,
            byte_size=size,
            checksum=checksum,
            fingerprint=fingerprint,
            artifact_ref=_safe_ref("artifact", fingerprint),
            content_type=DOCX_CONTENT_TYPE,
        )

    def _readiness(
        self, pilot: PilotStateRecord, checklist: DeliveryChecklist
    ) -> ReadinessView:
        items = self._item_records(checklist.checklist_id)
        blocking: list[ReadinessIssue] = []
        warnings: list[ReadinessIssue] = []
        fingerprints: list[str] = []
        stale = False
        if pilot.status != "approved":
            blocking.append(ReadinessIssue(code="pilot_not_approved"))
        current_template = self._current_template()
        if checklist.template_version != current_template.template_version:
            stale = True
            blocking.append(ReadinessIssue(code="template_version_stale"))
        for item in items:
            effective = ChecklistItemStatus(item.status)
            current_artifact: _ResolvedArtifact | None = None
            if effective == ChecklistItemStatus.PRESENT:
                current_artifact, current_status = self._resolve_item(pilot, item)
                if (
                    current_status != ChecklistItemStatus.PRESENT
                    or current_artifact is None
                    or current_artifact.fingerprint != item.artifact_fingerprint
                ):
                    effective = ChecklistItemStatus.STALE
                else:
                    fingerprints.append(
                        f"{item.item_key}:{current_artifact.fingerprint}"
                    )
            if effective == ChecklistItemStatus.STALE:
                stale = True
                blocking.append(
                    ReadinessIssue(item_key=item.item_key, code="artifact_stale")
                )
            elif effective == ChecklistItemStatus.WAIVED:
                if not item.waiver_allowed:
                    blocking.append(
                        ReadinessIssue(
                            item_key=item.item_key, code="waiver_not_allowed"
                        )
                    )
                else:
                    warnings.append(
                        ReadinessIssue(item_key=item.item_key, code="required_waived")
                    )
            elif effective != ChecklistItemStatus.PRESENT:
                code = f"artifact_{effective.value}"
                if item.requirement == "required":
                    blocking.append(ReadinessIssue(item_key=item.item_key, code=code))
                else:
                    warnings.append(ReadinessIssue(item_key=item.item_key, code=code))
        digest = _sha256_bytes("\n".join(sorted(fingerprints)).encode())
        status = (
            ReadinessStatus.STALE
            if stale
            else ReadinessStatus.READY if not blocking else ReadinessStatus.NOT_READY
        )
        return ReadinessView(
            pilot_ref=_safe_ref("pilot", pilot.pilot_id),
            pilot_status=pilot.status,
            pilot_version=pilot.version,
            checklist_id=checklist.checklist_id,
            checklist_version=checklist.version,
            template_version=checklist.template_version,
            readiness=status,
            blocking_items=blocking,
            warnings=warnings,
            artifact_set_fingerprint=digest,
            evaluated_at=self._now(),
        )

    def _checklist_view(
        self, checklist: DeliveryChecklist, *, derive_stale: bool = False
    ) -> ChecklistView:
        pilot = self._get_pilot(checklist.pilot_id)
        items = []
        for record in self._item_records(checklist.checklist_id):
            status = ChecklistItemStatus(record.status)
            artifact = None
            if status == ChecklistItemStatus.PRESENT:
                artifact, resolved_status = self._resolve_item(pilot, record)
                if derive_stale and (
                    resolved_status != ChecklistItemStatus.PRESENT
                    or artifact is None
                    or artifact.fingerprint != record.artifact_fingerprint
                ):
                    status = ChecklistItemStatus.STALE
            items.append(
                ChecklistItemView(
                    item_key=record.item_key,
                    display_name=record.display_name,
                    requirement=record.requirement,
                    artifact_type=record.artifact_type,
                    status=status,
                    artifact_ref=record.artifact_ref,
                    artifact_fingerprint=record.artifact_fingerprint,
                    byte_size=artifact.byte_size if artifact else None,
                    checksum=artifact.checksum if artifact else None,
                    content_type=artifact.content_type if artifact else None,
                    verified_at=_optional_utc(record.verified_at),
                    waived=status == ChecklistItemStatus.WAIVED,
                    version=record.version,
                )
            )
        return ChecklistView(
            checklist_id=checklist.checklist_id,
            pilot_ref=_safe_ref("pilot", checklist.pilot_id),
            template_version=checklist.template_version,
            version=checklist.version,
            created_at=_as_utc(checklist.created_at),
            updated_at=_as_utc(checklist.updated_at),
            items=items,
        )

    def _assembly_view(self, assembly: DeliveryPackageAssembly) -> AssemblyView:
        package = (
            self.db.query(DeliveryPackage)
            .filter(DeliveryPackage.assembly_id == assembly.assembly_id)
            .first()
        )
        return AssemblyView(
            assembly_id=assembly.assembly_id,
            pilot_ref=_safe_ref("pilot", assembly.pilot_id),
            status=AssemblyStatus(assembly.status),
            version=assembly.version,
            attempt=assembly.attempt,
            failure_code=assembly.failure_code,
            package_ref=_safe_ref("package", package.package_id) if package else None,
            created_at=_as_utc(assembly.created_at),
            updated_at=_as_utc(assembly.updated_at),
            completed_at=_optional_utc(assembly.completed_at),
        )

    def _package_summary(self, package: DeliveryPackage) -> PackageSummary:
        return PackageSummary(
            package_id=package.package_id,
            package_ref=_safe_ref("package", package.package_id),
            status=package.status,
            manifest_version=package.manifest_version,
            byte_size=package.byte_size,
            checksum=package.checksum,
            created_at=_as_utc(package.created_at),
        )

    def _ensure_template(self) -> DeliveryChecklistTemplate:
        existing = (
            self.db.query(DeliveryChecklistTemplate)
            .filter(
                DeliveryChecklistTemplate.template_key == DEFAULT_TEMPLATE_KEY,
                DeliveryChecklistTemplate.template_version == DEFAULT_TEMPLATE_VERSION,
            )
            .first()
        )
        if existing is not None:
            return existing
        now = self._now()
        template = DeliveryChecklistTemplate(
            template_id=uuid.uuid4().hex,
            template_key=DEFAULT_TEMPLATE_KEY,
            template_version=DEFAULT_TEMPLATE_VERSION,
            name="Pilot delivery checklist v1",
            created_at=now,
        )
        self.db.add(template)
        self.db.flush()
        for index, spec in enumerate(DEFAULT_TEMPLATE_ITEMS):
            self.db.add(
                DeliveryChecklistTemplateItem(
                    template_item_id=uuid.uuid4().hex,
                    template_id=template.template_id,
                    item_key=spec.item_key,
                    display_name=spec.display_name,
                    description=spec.description,
                    requirement=spec.requirement,
                    artifact_type=spec.artifact_type,
                    source=spec.source,
                    waiver_allowed=spec.waiver_allowed,
                    sort_order=index,
                )
            )
        self.db.flush()
        return template

    def _current_template(self) -> DeliveryChecklistTemplate:
        template = (
            self.db.query(DeliveryChecklistTemplate)
            .filter(
                DeliveryChecklistTemplate.template_key == DEFAULT_TEMPLATE_KEY,
                DeliveryChecklistTemplate.template_version == DEFAULT_TEMPLATE_VERSION,
            )
            .first()
        )
        if template is None:
            raise DeliveryNotFoundError("Delivery checklist template was not found")
        return template

    def _pilot_and_project(
        self, pilot_id: str
    ) -> tuple[PilotStateRecord, ModernizationProject]:
        pilot = self._get_pilot(pilot_id)
        return pilot, self._get_project(pilot.project_id)

    def _get_pilot(self, pilot_id: str) -> PilotStateRecord:
        record = (
            self.db.query(PilotStateRecord)
            .filter(PilotStateRecord.pilot_id == pilot_id)
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Pilot was not found")
        return record

    def _get_project(self, project_id: str) -> ModernizationProject:
        record = (
            self.db.query(ModernizationProject)
            .filter(ModernizationProject.id == project_id)
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Project was not found")
        return record

    def _find_checklist(
        self, pilot_id: str, template_id: str
    ) -> DeliveryChecklist | None:
        return (
            self.db.query(DeliveryChecklist)
            .filter(
                DeliveryChecklist.pilot_id == pilot_id,
                DeliveryChecklist.template_id == template_id,
                DeliveryChecklist.template_version == DEFAULT_TEMPLATE_VERSION,
            )
            .first()
        )

    def _get_checklist(self, pilot_id: str) -> DeliveryChecklist:
        record = (
            self.db.query(DeliveryChecklist)
            .filter(DeliveryChecklist.pilot_id == pilot_id)
            .order_by(DeliveryChecklist.template_version.desc())
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Delivery checklist was not found")
        return record

    def _get_item_record(
        self, checklist_id: str, item_key: str
    ) -> DeliveryChecklistItem:
        record = (
            self.db.query(DeliveryChecklistItem)
            .filter(
                DeliveryChecklistItem.checklist_id == checklist_id,
                DeliveryChecklistItem.item_key == item_key,
            )
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Checklist item was not found")
        return record

    def _item_records(self, checklist_id: str) -> list[DeliveryChecklistItem]:
        return (
            self.db.query(DeliveryChecklistItem)
            .filter(DeliveryChecklistItem.checklist_id == checklist_id)
            .order_by(DeliveryChecklistItem.sort_order.asc())
            .all()
        )

    def _get_assembly(self, assembly_id: str) -> DeliveryPackageAssembly:
        record = (
            self.db.query(DeliveryPackageAssembly)
            .filter(DeliveryPackageAssembly.assembly_id == assembly_id)
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Assembly was not found")
        return record

    def _get_package(self, package_id: str) -> DeliveryPackage:
        record = (
            self.db.query(DeliveryPackage)
            .filter(DeliveryPackage.package_id == package_id)
            .first()
        )
        if record is None:
            raise DeliveryNotFoundError("Package was not found")
        return record

    def _require_project_access(
        self, actor: User, project: ModernizationProject
    ) -> None:
        if actor.role == UserRole.ADMIN.value or (
            actor.role == UserRole.USER.value and project.user_id == actor.id
        ):
            return
        raise DeliveryAccessDeniedError("Delivery access is denied")

    def _require_operator(self, actor: User) -> None:
        if actor.role != UserRole.ADMIN.value:
            raise DeliveryAccessDeniedError("Operator capability is required")

    def _require_mutable(self, pilot: PilotStateRecord) -> None:
        if pilot.status == "delivered":
            raise DeliveryStateConflictError("Delivered Pilot is read-only")

    def _require_not_frozen(self, checklist_id: str) -> None:
        frozen = (
            self.db.query(DeliveryPackageAssembly.assembly_id)
            .filter(
                DeliveryPackageAssembly.checklist_id == checklist_id,
                DeliveryPackageAssembly.status == AssemblyStatus.ASSEMBLING.value,
            )
            .first()
        )
        if frozen is not None:
            raise DeliveryStateConflictError("Checklist is frozen during assembly")

    def _require_checklist_version(
        self, checklist: DeliveryChecklist, expected_version: int
    ) -> None:
        if checklist.version != expected_version:
            raise DeliveryVersionConflictError(checklist.version)

    def _load_replay(
        self,
        actor_id: int,
        idempotency_key: str,
        operation: str,
        request_hash: str,
        model: type[BaseModel],
    ) -> BaseModel | None:
        record = (
            self.db.query(DeliveryCommandResult)
            .filter(
                DeliveryCommandResult.actor_id == actor_id,
                DeliveryCommandResult.idempotency_key == idempotency_key,
            )
            .first()
        )
        if record is None:
            return None
        if record.operation != operation or record.request_hash != request_hash:
            raise DeliveryIdempotencyConflictError(
                "Idempotency key was reused for a different command"
            )
        return model.model_validate_json(record.response_json)

    def _store_result(
        self,
        actor_id: int,
        idempotency_key: str,
        operation: str,
        request_hash: str,
        resource_type: str,
        resource_id: str,
        result_version: int,
        response: BaseModel,
        created_at: datetime,
    ) -> None:
        self.db.add(
            DeliveryCommandResult(
                actor_id=actor_id,
                idempotency_key=idempotency_key,
                operation=operation,
                request_hash=request_hash,
                resource_type=resource_type,
                resource_id=resource_id,
                result_version=result_version,
                response_json=response.model_dump_json(),
                created_at=created_at,
            )
        )

    def _replace_command_response(
        self, idempotency_key: str, actor_id: int, response: AssemblyView
    ) -> None:
        record = (
            self.db.query(DeliveryCommandResult)
            .filter(
                DeliveryCommandResult.actor_id == actor_id,
                DeliveryCommandResult.idempotency_key == idempotency_key,
            )
            .first()
        )
        if record is not None:
            record.result_version = response.version
            record.response_json = response.model_dump_json()

    def _add_audit(
        self,
        pilot: PilotStateRecord,
        actor: User,
        event_type: str,
        occurred_at: datetime,
        *,
        checklist_id: str | None = None,
        assembly_id: str | None = None,
        package_id: str | None = None,
        idempotency_key: str | None = None,
        result_version: int,
    ) -> None:
        self.db.add(
            DeliveryAuditEvent(
                event_id=uuid.uuid4().hex,
                pilot_id=pilot.pilot_id,
                project_id=pilot.project_id,
                run_id=pilot.run_id,
                checklist_id=checklist_id,
                assembly_id=assembly_id,
                package_id=package_id,
                event_type=event_type,
                actor_id=actor.id,
                occurred_at=occurred_at,
                idempotency_key=idempotency_key,
                result_version=result_version,
                metadata_json="{}",
            )
        )

    def _verify_package_file(self, package: DeliveryPackage) -> Path:
        paths = build_project_result_archive_paths(
            archive_root=self.archive_root,
            project_id=package.project_id,
            run_id=package.run_id,
        )
        self._assert_safe_root(paths["dir"], self.archive_root)
        path = paths["dir"] / "delivery_packages" / f"{package.artifact_reference}.zip"
        self._assert_safe_root(path, paths["dir"] / "delivery_packages")
        if not path.is_file() or path.is_symlink():
            raise PackageIntegrityError("Package file is unavailable")
        if (
            path.stat().st_size != package.byte_size
            or _sha256_path(path) != package.checksum
        ):
            raise PackageIntegrityError("Package integrity verification failed")
        return path

    def _assert_safe_root(self, path: Path, root: Path) -> None:
        resolved_root = root.resolve()
        resolved_path = path.resolve()
        try:
            resolved_path.relative_to(resolved_root)
        except ValueError as exc:
            raise InvalidArtifactError("Artifact is outside the allowed root") from exc

    def _cleanup_staging(self, now: datetime) -> None:
        if not self.archive_root.exists():
            return
        cutoff = now.timestamp() - STAGING_RETENTION.total_seconds()
        for staging in self.archive_root.glob("*/*/delivery_packages/.staging"):
            if staging.is_symlink() or not staging.is_dir():
                continue
            for candidate in staging.iterdir():
                try:
                    if (
                        candidate.is_file()
                        and not candidate.is_symlink()
                        and candidate.stat().st_mtime < cutoff
                    ):
                        candidate.unlink()
                except OSError:
                    continue

    def _now(self) -> datetime:
        value = self.now_provider()
        if value.tzinfo is None or value.utcoffset() is None:
            raise ValueError("now_provider must return timezone-aware datetime")
        return value.astimezone(timezone.utc)


def _validate_docx(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise InvalidArtifactError("DOCX structure is invalid")
        Document(str(path))
    except (OSError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise InvalidArtifactError("DOCX structure is invalid") from exc


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _write_zip_entry(archive: zipfile.ZipFile, name: str, content: bytes) -> None:
    info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
    info.compress_type = zipfile.ZIP_DEFLATED
    info.external_attr = 0o100644 << 16
    archive.writestr(info, content, compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def _safe_failure_code(exc: Exception) -> str:
    if isinstance(exc, InvalidArtifactError):
        message = str(exc).lower()
        return "package_size_exceeded" if "size" in message else "invalid_artifact"
    if isinstance(exc, PackageIntegrityError):
        return "integrity_check_failed"
    if isinstance(exc, (OSError, SQLAlchemyError, DeliveryStorageError)):
        return "assembly_storage_error"
    return "assembly_failed"


def _request_hash(operation: str, payload: dict[str, object]) -> str:
    canonical = json.dumps(
        {"operation": operation, "payload": payload},
        ensure_ascii=True,
        sort_keys=True,
        separators=(",", ":"),
    )
    return _sha256_bytes(canonical.encode())


def _safe_ref(prefix: str, value: str) -> str:
    return f"{prefix}-{_sha256_bytes(value.encode())[:12]}"


def _sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _sha256_path(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _as_utc(value: datetime) -> datetime:
    if value.tzinfo is None or value.utcoffset() is None:
        return value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc)


def _optional_utc(value: datetime | None) -> datetime | None:
    return _as_utc(value) if value is not None else None
