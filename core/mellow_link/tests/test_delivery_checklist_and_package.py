from __future__ import annotations

import hashlib
import inspect as python_inspect
import json
import os
import shutil
import threading
import zipfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from pathlib import Path

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event, inspect
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import sessionmaker

import mellow_link.routers.delivery_package as delivery_router_module
import mellow_link.services.delivery_package as delivery_module
from mellow_link.infra.database import (
    AgentRun,
    Base,
    DeliveryAuditEvent,
    DeliveryChecklist,
    DeliveryCommandResult,
    DeliveryDownloadReference,
    DeliveryPackage,
    DeliveryPackageAssembly,
    ModernizationProject,
    PilotStateRecord,
    ProjectAsset,
    User,
    UserRole,
    get_current_user,
    get_db,
)
from mellow_link.routers.delivery_package import (
    request_assembly as request_assembly_route,
)
from mellow_link.routers.delivery_package import router
from mellow_link.services.delivery_package import (
    AssemblyStatus,
    ChecklistItemStatus,
    DeliveryAccessDeniedError,
    DeliveryIdempotencyConflictError,
    DeliveryPackageService,
    DeliveryStateConflictError,
    DeliveryStorageError,
    DeliveryValidationError,
    DeliveryVersionConflictError,
    IdempotentRequest,
    PackageIntegrityError,
    ReadinessBlockedError,
    ReadinessStaleError,
    ReadinessStatus,
    RequestAssemblyRequest,
    RetryAssemblyRequest,
    VersionedChecklistRequest,
    WaiveChecklistItemRequest,
)
from mellow_link.services.project_results.archive import (
    build_project_result_archive_paths,
)


class Clock:
    def __init__(self) -> None:
        self.value = datetime(2026, 7, 20, 0, 0, tzinfo=timezone.utc)

    def __call__(self) -> datetime:
        value = self.value
        self.value += timedelta(seconds=1)
        return value


@pytest.fixture()
def database(tmp_path):
    engine = create_engine(
        f"sqlite:///{tmp_path / 'delivery.db'}",
        connect_args={"check_same_thread": False},
    )

    @event.listens_for(engine, "connect")
    def enable_foreign_keys(connection, _record):
        connection.execute("PRAGMA foreign_keys=ON")

    Base.metadata.create_all(bind=engine)
    factory = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    try:
        yield engine, factory
    finally:
        engine.dispose()


@pytest.fixture()
def db(database):
    _, factory = database
    session = factory()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture()
def actors(db):
    owner = User(
        username="delivery-owner",
        hashed_password="synthetic-owner-hash",
        role=UserRole.USER.value,
    )
    other = User(
        username="delivery-other",
        hashed_password="synthetic-other-hash",
        role=UserRole.USER.value,
    )
    operator = User(
        username="delivery-operator",
        hashed_password="synthetic-operator-hash",
        role=UserRole.ADMIN.value,
    )
    operator_two = User(
        username="delivery-operator-two",
        hashed_password="synthetic-operator-two-hash",
        role=UserRole.ADMIN.value,
    )
    guest = User(
        username="delivery-guest",
        hashed_password="synthetic-guest-hash",
        role=UserRole.GUEST.value,
    )
    db.add_all([owner, other, operator, operator_two, guest])
    db.commit()
    for actor in (owner, other, operator, operator_two, guest):
        db.refresh(actor)
    return owner, other, operator, operator_two, guest


@pytest.fixture()
def pilot(db, actors):
    owner, *_ = actors
    run = AgentRun(
        run_id="delivery-run",
        session_id="delivery-session",
        module_id="engine",
        run_kind="project",
        status="completed",
    )
    project = ModernizationProject(
        id="delivery-project",
        user_id=owner.id,
        session_id="delivery-session",
        run_id=run.run_id,
        project_name="Synthetic Delivery Project",
        client_name="Synthetic Client",
        goal_text="Synthetic delivery verification",
        upload_session_id="synthetic-upload",
        status="completed",
    )
    state = PilotStateRecord(
        pilot_id="delivery-pilot",
        project_id=project.id,
        run_id=run.run_id,
        status="approved",
        version=3,
        created_at=datetime(2026, 7, 20, tzinfo=timezone.utc),
        updated_at=datetime(2026, 7, 20, tzinfo=timezone.utc),
        created_by_id=owner.id,
        approved_by_id=actors[2].id,
        approved_at=datetime(2026, 7, 20, tzinfo=timezone.utc),
    )
    db.add_all([run, project, state])
    db.commit()
    return state


@pytest.fixture()
def clock():
    return Clock()


@pytest.fixture()
def service(db, tmp_path, clock):
    return DeliveryPackageService(
        db,
        archive_root=tmp_path / "project-results",
        now_provider=clock,
    )


def write_external_docx(service: DeliveryPackageService, pilot, *, internal=False):
    path = build_project_result_archive_paths(
        archive_root=service.archive_root,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
    )["external_docx"]
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for index in range(1, 10):
        if index == 1:
            title = "1. 1페이지 요약"
        elif index == 8:
            title = "8. 단계별 실행 준비 계획"
        elif index == 9:
            title = "9. 분석 근거와 provenance" if internal else "9. 산출물 기준"
        else:
            title = f"{index}. Synthetic section"
        document.add_heading(title, level=1)
        document.add_paragraph(f"Synthetic content for section {index}.")
    if internal:
        document.add_paragraph("run_id: internal-run")
    document.save(path)
    return path


def create_checklist(service, operator, pilot, key="create-checklist"):
    return service.create_checklist(
        operator, pilot.pilot_id, IdempotentRequest(idempotency_key=key)
    )


def verify_required_items(service, operator, pilot, checklist):
    current = checklist
    for item_key in (
        "external_report_docx",
        "one_page_summary",
        "provenance_summary",
    ):
        current = service.verify_item(
            operator,
            pilot.pilot_id,
            item_key,
            VersionedChecklistRequest(
                expected_checklist_version=current.version,
                idempotency_key=f"verify-{item_key}",
            ),
        )
    return current


def ready_checklist(service, operator, pilot):
    write_external_docx(service, pilot)
    checklist = create_checklist(service, operator, pilot)
    checklist = verify_required_items(service, operator, pilot, checklist)
    readiness = service.get_readiness(operator, pilot.pilot_id)
    assert readiness.readiness == ReadinessStatus.READY
    return checklist, readiness


def assemble(service, operator, pilot):
    checklist, readiness = ready_checklist(service, operator, pilot)
    result = service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="assemble-package",
        ),
    )
    assert result.status == AssemblyStatus.ASSEMBLED
    return result


def test_additive_schema_contains_delivery_tables(database):
    engine, factory = database
    tables = set(inspect(engine).get_table_names())
    assert {
        "delivery_checklist_templates",
        "delivery_checklist_template_items",
        "delivery_checklists",
        "delivery_checklist_items",
        "delivery_package_assemblies",
        "delivery_packages",
        "delivery_audit_events",
        "delivery_command_results",
        "delivery_download_references",
    }.issubset(tables)
    session = factory()
    try:
        existing = User(
            username="existing-schema-user",
            hashed_password="synthetic-schema-hash",
            role=UserRole.USER.value,
        )
        session.add(existing)
        session.commit()
        Base.metadata.create_all(bind=engine)
        assert session.query(User).filter(User.username == existing.username).one()
    finally:
        session.close()


def test_create_checklist_is_persistent_idempotent_and_versioned(
    database, db, actors, pilot, service, tmp_path, clock
):
    _, _, operator, *_ = actors
    created = create_checklist(service, operator, pilot)
    replay = create_checklist(service, operator, pilot)

    assert replay == created
    assert created.version == 0
    assert [item.item_key for item in created.items] == [
        "external_report_docx",
        "one_page_summary",
        "provenance_summary",
        "delivery_note",
    ]
    assert db.query(DeliveryChecklist).count() == 1
    assert db.query(DeliveryAuditEvent).count() == 1
    assert db.query(DeliveryCommandResult).count() == 1

    _, factory = database
    db.close()
    restarted = factory()
    try:
        restarted_operator = restarted.query(User).filter(User.id == operator.id).one()
        restarted_service = DeliveryPackageService(
            restarted,
            archive_root=tmp_path / "project-results",
            now_provider=clock,
        )
        assert (
            restarted_service.get_checklist(restarted_operator, pilot.pilot_id)
            == created
        )
    finally:
        restarted.close()


def test_checklist_permissions_and_project_isolation(actors, pilot, service):
    owner, other, operator, _, guest = actors
    with pytest.raises(DeliveryAccessDeniedError):
        create_checklist(service, owner, pilot)
    created = create_checklist(service, operator, pilot)
    assert service.get_checklist(owner, pilot.pilot_id) == created
    for actor in (other, guest):
        with pytest.raises(DeliveryAccessDeniedError):
            service.get_checklist(actor, pilot.pilot_id)


def test_verify_required_items_and_derive_ready_with_optional_warning(
    actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)

    assert checklist.version == 3
    assert all(
        item.status == ChecklistItemStatus.PRESENT
        for item in checklist.items
        if item.requirement == "required"
    )
    assert all(
        item.verified_by_ref and item.verified_by_ref != str(operator.id)
        for item in checklist.items
        if item.requirement == "required"
    )
    assert readiness.readiness == ReadinessStatus.READY
    assert [warning.code for warning in readiness.warnings] == ["artifact_pending"]
    assert len(readiness.artifact_set_fingerprint) == 64


def test_missing_and_internal_docx_are_not_ready(actors, pilot, service):
    operator = actors[2]
    checklist = create_checklist(service, operator, pilot)
    checklist = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=0,
            idempotency_key="verify-missing",
        ),
    )
    assert checklist.items[0].status == ChecklistItemStatus.MISSING
    assert (
        service.get_readiness(operator, pilot.pilot_id).readiness
        == ReadinessStatus.NOT_READY
    )

    write_external_docx(service, pilot, internal=True)
    checklist = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=1,
            idempotency_key="verify-internal",
        ),
    )
    assert checklist.items[0].status == ChecklistItemStatus.INVALID


def test_oversized_external_docx_is_invalid(monkeypatch, actors, pilot, service):
    operator = actors[2]
    path = write_external_docx(service, pilot)
    assert path.stat().st_size > 128
    monkeypatch.setattr(delivery_module, "MAX_ARTIFACT_BYTES", 128)
    checklist = create_checklist(service, operator, pilot)
    verified = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="verify-oversized",
        ),
    )
    item = next(
        item for item in verified.items if item.item_key == "external_report_docx"
    )
    assert item.status == ChecklistItemStatus.INVALID


def test_package_output_size_limit_fails_safely(
    monkeypatch, db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    monkeypatch.setattr(delivery_module, "MAX_PACKAGE_BYTES", 1)
    result = service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="package-size-limit",
        ),
    )
    assert result.status == AssemblyStatus.FAILED
    assert result.failure_code == "package_size_exceeded"
    assert db.query(DeliveryPackage).count() == 0
    assert not list(service.archive_root.rglob("*.zip"))


def test_symbolic_link_artifact_is_rejected(tmp_path, actors, pilot, service):
    operator = actors[2]
    outside = tmp_path / "outside.docx"
    document = Document()
    for index in range(1, 10):
        document.add_heading(
            (
                "1. 1페이지 요약"
                if index == 1
                else "9. 산출물 기준" if index == 9 else f"{index}. Synthetic section"
            ),
            level=1,
        )
    document.save(outside)
    path = build_project_result_archive_paths(
        archive_root=service.archive_root,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
    )["external_docx"]
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        path.symlink_to(outside)
    except OSError:
        pytest.skip("symbolic link creation is unavailable")
    checklist = create_checklist(service, operator, pilot)
    verified = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="verify-symlink",
        ),
    )
    item = next(
        item for item in verified.items if item.item_key == "external_report_docx"
    )
    assert item.status == ChecklistItemStatus.INVALID


def test_hard_link_artifact_is_rejected(tmp_path, actors, pilot, service):
    operator = actors[2]
    path = write_external_docx(service, pilot)
    outside = tmp_path / "hard-linked.docx"
    try:
        os.link(path, outside)
    except OSError:
        pytest.skip("hard link creation is unavailable")
    checklist = create_checklist(service, operator, pilot)
    verified = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="verify-hard-link",
        ),
    )
    item = next(
        item for item in verified.items if item.item_key == "external_report_docx"
    )
    assert item.status == ChecklistItemStatus.INVALID


def test_source_filename_in_external_docx_is_rejected(db, actors, pilot, service):
    operator = actors[2]
    path = write_external_docx(service, pilot)
    db.add(
        ProjectAsset(
            id="synthetic-asset",
            project_id=pilot.project_id,
            source_temp_session_id="synthetic-session",
            source_temp_file_id="synthetic-file",
            original_filename="sensitive-source.sql",
            stored_relative_path="safe/source",
            extracted_relative_path="safe/extracted",
            file_size=32,
            content_type="text/plain",
            category_hint="source",
            extracted_chars=0,
        )
    )
    db.commit()
    document = Document(path)
    document.add_paragraph("sensitive-source.sql")
    document.save(path)
    checklist = create_checklist(service, operator, pilot)
    verified = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="verify-source-name",
        ),
    )
    item = next(
        item for item in verified.items if item.item_key == "external_report_docx"
    )
    assert item.status == ChecklistItemStatus.INVALID


def test_changed_artifact_is_stale_and_blocks_assembly(actors, pilot, service):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    path = build_project_result_archive_paths(
        archive_root=service.archive_root,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
    )["external_docx"]
    document = Document(path)
    document.add_paragraph("Changed synthetic content")
    document.save(path)

    current = service.get_checklist(operator, pilot.pilot_id)
    assert current.items[0].status == ChecklistItemStatus.STALE
    assert (
        service.get_readiness(operator, pilot.pilot_id).readiness
        == ReadinessStatus.STALE
    )
    with pytest.raises((ReadinessStaleError, ReadinessBlockedError)) as captured:
        service.request_assembly(
            operator,
            pilot.pilot_id,
            RequestAssemblyRequest(
                expected_pilot_version=pilot.version,
                expected_checklist_version=checklist.version,
                artifact_set_fingerprint=readiness.artifact_set_fingerprint,
                idempotency_key="stale-assembly",
            ),
        )
    assert getattr(captured.value, "code", "") in {
        "readiness_stale",
        "readiness_blocked",
    }


def test_version_and_idempotency_conflicts_leave_state_unchanged(
    db, actors, pilot, service
):
    operator = actors[2]
    write_external_docx(service, pilot)
    checklist = create_checklist(service, operator, pilot)
    with pytest.raises(DeliveryVersionConflictError):
        service.verify_item(
            operator,
            pilot.pilot_id,
            "external_report_docx",
            VersionedChecklistRequest(
                expected_checklist_version=8,
                idempotency_key="stale-version",
            ),
        )
    first = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="same-key",
        ),
    )
    replay = service.verify_item(
        operator,
        pilot.pilot_id,
        "external_report_docx",
        VersionedChecklistRequest(
            expected_checklist_version=checklist.version,
            idempotency_key="same-key",
        ),
    )
    assert replay == first
    with pytest.raises(DeliveryIdempotencyConflictError):
        service.verify_item(
            operator,
            pilot.pilot_id,
            "one_page_summary",
            VersionedChecklistRequest(
                expected_checklist_version=first.version,
                idempotency_key="same-key",
            ),
        )
    assert db.query(DeliveryAuditEvent).count() == 2


def test_initial_required_items_cannot_be_waived(actors, pilot, service):
    operator = actors[2]
    checklist = create_checklist(service, operator, pilot)
    with pytest.raises(DeliveryStateConflictError):
        service.waive_item(
            operator,
            pilot.pilot_id,
            "external_report_docx",
            WaiveChecklistItemRequest(
                expected_checklist_version=checklist.version,
                idempotency_key="waive-report",
                reason="Synthetic waiver reason",
            ),
        )


def test_checklist_audit_failure_rolls_back_item_and_version(
    monkeypatch, db, actors, pilot, service
):
    operator = actors[2]
    write_external_docx(service, pilot)
    create_checklist(service, operator, pilot)

    def fail_audit(*args, **kwargs):
        raise SQLAlchemyError("synthetic audit failure")

    monkeypatch.setattr(service, "_add_audit", fail_audit)
    with pytest.raises(Exception):
        service.verify_item(
            operator,
            pilot.pilot_id,
            "external_report_docx",
            VersionedChecklistRequest(
                expected_checklist_version=0,
                idempotency_key="audit-failure",
            ),
        )
    db.expire_all()
    stored = db.query(DeliveryChecklist).one()
    assert stored.version == 0
    assert (
        service.get_item(operator, pilot.pilot_id, "external_report_docx").status
        == ChecklistItemStatus.PENDING
    )


def test_assembly_creates_reopenable_deterministic_safe_package(
    db, actors, pilot, service
):
    operator = actors[2]
    result = assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    path = (
        build_project_result_archive_paths(
            archive_root=service.archive_root,
            project_id=pilot.project_id,
            run_id=pilot.run_id,
        )["dir"]
        / "delivery_packages"
        / f"{package.artifact_reference}.zip"
    )

    assert path.is_file()
    assert package.checksum == hashlib.sha256(path.read_bytes()).hexdigest()
    with zipfile.ZipFile(path) as archive:
        assert archive.namelist() == ["manifest.json", "report/pilot-report.docx"]
        manifest = json.loads(archive.read("manifest.json"))
        Document(archive.open("report/pilot-report.docx"))
    serialized = json.dumps(manifest)
    assert manifest["assembly_status"] == "assembled"
    assert manifest["logical_coverage"] == ["one_page_summary", "external_provenance"]
    assert str(service.archive_root) not in serialized
    assert "external_result.docx" not in serialized
    assert "delivery-run" not in serialized
    assert db.query(DeliveryPackageAssembly).one().version == 2
    assert result.package_ref is not None
    artifact_dir = os.environ.get("PILOT_DELIVERY_ARTIFACT_DIR")
    if artifact_dir:
        destination = Path(artifact_dir) / "synthetic-pilot-delivery-package.zip"
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(path, destination)


def test_assembly_idempotency_replays_without_duplicate_package_or_audit(
    db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    payload = RequestAssemblyRequest(
        expected_pilot_version=pilot.version,
        expected_checklist_version=checklist.version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        idempotency_key="assembly-replay",
    )
    first = service.request_assembly(operator, pilot.pilot_id, payload)
    count = db.query(DeliveryAuditEvent).count()
    replay = service.request_assembly(operator, pilot.pilot_id, payload)
    assert replay == first
    assert db.query(DeliveryPackage).count() == 1
    assert db.query(DeliveryAuditEvent).count() == count


def test_assembly_idempotency_key_with_different_payload_is_rejected(
    actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    payload = RequestAssemblyRequest(
        expected_pilot_version=pilot.version,
        expected_checklist_version=checklist.version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        idempotency_key="assembly-conflict-key",
    )
    service.request_assembly(operator, pilot.pilot_id, payload)
    with pytest.raises(DeliveryIdempotencyConflictError):
        service.request_assembly(
            operator,
            pilot.pilot_id,
            RequestAssemblyRequest(
                expected_pilot_version=pilot.version,
                expected_checklist_version=checklist.version,
                artifact_set_fingerprint="0" * 64,
                idempotency_key="assembly-conflict-key",
            ),
        )


def test_assembly_idempotency_survives_service_restart(
    database, db, actors, pilot, service, tmp_path, clock
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    payload = RequestAssemblyRequest(
        expected_pilot_version=pilot.version,
        expected_checklist_version=checklist.version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        idempotency_key="assembly-restart-replay",
    )
    first = service.request_assembly(operator, pilot.pilot_id, payload)
    _, factory = database
    restarted_db = factory()
    try:
        restarted_operator = (
            restarted_db.query(User).filter(User.id == operator.id).one()
        )
        restarted = DeliveryPackageService(
            restarted_db,
            archive_root=tmp_path / "project-results",
            now_provider=clock,
        )
        replay = restarted.request_assembly(restarted_operator, pilot.pilot_id, payload)
        assert replay == first
        assert restarted_db.query(DeliveryPackage).count() == 1
    finally:
        restarted_db.close()


def test_concurrent_identical_assembly_request_creates_one_package(
    database, db, actors, pilot, service, tmp_path
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    payload = RequestAssemblyRequest(
        expected_pilot_version=pilot.version,
        expected_checklist_version=checklist.version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        idempotency_key="concurrent-identical-assembly",
    )
    _, factory = database
    barrier = threading.Barrier(2)
    now = datetime(2026, 7, 20, 0, 10, tzinfo=timezone.utc)

    def request() -> AssemblyStatus:
        session = factory()
        try:
            concurrent_operator = (
                session.query(User).filter(User.id == operator.id).one()
            )
            concurrent_service = DeliveryPackageService(
                session,
                archive_root=tmp_path / "project-results",
                now_provider=lambda: now,
            )
            barrier.wait(timeout=5)
            return concurrent_service.request_assembly(
                concurrent_operator, pilot.pilot_id, payload
            ).status
        finally:
            session.close()

    with ThreadPoolExecutor(max_workers=2) as pool:
        statuses = list(pool.map(lambda _index: request(), range(2)))
    assert set(statuses) <= {
        AssemblyStatus.PENDING,
        AssemblyStatus.ASSEMBLING,
        AssemblyStatus.ASSEMBLED,
    }
    db.expire_all()
    assert db.query(DeliveryPackageAssembly).count() == 1
    assert db.query(DeliveryPackage).count() == 1
    assert (
        db.query(DeliveryCommandResult)
        .filter(
            DeliveryCommandResult.idempotency_key == "concurrent-identical-assembly"
        )
        .count()
        == 1
    )


def test_storage_constraint_and_conditional_claim_prevent_duplicate_assembly(
    database, db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    assembly = DeliveryPackageAssembly(
        assembly_id="conditional-assembly",
        pilot_id=pilot.pilot_id,
        checklist_id=checklist.checklist_id,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
        status="pending",
        version=0,
        attempt=1,
        request_fingerprint="c" * 64,
        source_pilot_version=pilot.version,
        checklist_version=checklist.version,
        template_version=checklist.template_version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        manifest_version="delivery-manifest-v1",
        created_by_id=operator.id,
        created_at=service._now(),
        updated_at=service._now(),
    )
    db.add(assembly)
    db.commit()
    _, factory = database
    first_db = factory()
    second_db = factory()
    try:
        first = (
            first_db.query(DeliveryPackageAssembly)
            .filter(
                DeliveryPackageAssembly.assembly_id == assembly.assembly_id,
                DeliveryPackageAssembly.status == "pending",
                DeliveryPackageAssembly.version == 0,
            )
            .update({"status": "assembling", "version": 1})
        )
        first_db.commit()
        second = (
            second_db.query(DeliveryPackageAssembly)
            .filter(
                DeliveryPackageAssembly.assembly_id == assembly.assembly_id,
                DeliveryPackageAssembly.status == "pending",
                DeliveryPackageAssembly.version == 0,
            )
            .update({"status": "assembling", "version": 1})
        )
        second_db.commit()
        assert (first, second) == (1, 0)
    finally:
        first_db.close()
        second_db.close()


def test_assembly_requires_approved_pilot_and_current_snapshot(
    db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    pilot.status = "under_review"
    db.commit()
    with pytest.raises(DeliveryStateConflictError):
        service.request_assembly(
            operator,
            pilot.pilot_id,
            RequestAssemblyRequest(
                expected_pilot_version=pilot.version,
                expected_checklist_version=checklist.version,
                artifact_set_fingerprint=readiness.artifact_set_fingerprint,
                idempotency_key="not-approved",
            ),
        )


def test_package_access_and_integrity_checks(db, actors, pilot, service):
    owner, other, operator, *_ = actors
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    manifest = service.get_manifest(owner, package.package_id)
    assert manifest.manifest["manifest_version"] == "delivery-manifest-v1"
    with pytest.raises(DeliveryAccessDeniedError):
        service.get_manifest(other, package.package_id)
    path = service._verify_package_file(package)
    path.write_bytes(path.read_bytes() + b"tampered")
    with pytest.raises(PackageIntegrityError):
        service.get_manifest(operator, package.package_id)


def test_download_reference_is_digest_only_single_use_and_authorized(
    db, actors, pilot, service
):
    _, _, operator, operator_two, _ = actors
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    reference = service.create_download_reference(operator, package.package_id)
    stored = db.query(DeliveryDownloadReference).one()
    assert reference.download_reference not in stored.token_digest
    assert (
        stored.token_digest
        == hashlib.sha256(reference.download_reference.encode()).hexdigest()
    )
    with pytest.raises(DeliveryAccessDeniedError):
        service.redeem_download_reference(operator_two, reference.download_reference)
    resolved = service.redeem_download_reference(operator, reference.download_reference)
    assert resolved.path.is_file()
    assert resolved.filename == "pilot-delivery-package.zip"
    with pytest.raises(DeliveryStateConflictError):
        service.redeem_download_reference(operator, reference.download_reference)


def test_expired_download_reference_is_rejected(db, actors, pilot, service, clock):
    operator = actors[2]
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    reference = service.create_download_reference(operator, package.package_id)
    clock.value += timedelta(minutes=16)
    with pytest.raises(DeliveryStateConflictError):
        service.redeem_download_reference(operator, reference.download_reference)


def test_download_reference_expiry_boundary_is_utc_and_exclusive(
    db, actors, pilot, service, clock
):
    operator = actors[2]
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    before = service.create_download_reference(operator, package.package_id)
    before_record = (
        db.query(DeliveryDownloadReference)
        .filter(
            DeliveryDownloadReference.token_digest
            == hashlib.sha256(before.download_reference.encode()).hexdigest()
        )
        .one()
    )
    clock.value = before_record.expires_at.replace(tzinfo=timezone.utc) - timedelta(
        microseconds=1
    )
    assert service.redeem_download_reference(operator, before.download_reference)

    exact = service.create_download_reference(operator, package.package_id)
    exact_record = (
        db.query(DeliveryDownloadReference)
        .filter(
            DeliveryDownloadReference.token_digest
            == hashlib.sha256(exact.download_reference.encode()).hexdigest()
        )
        .one()
    )
    clock.value = exact_record.expires_at.replace(tzinfo=timezone.utc)
    with pytest.raises(DeliveryStateConflictError):
        service.redeem_download_reference(operator, exact.download_reference)


def test_delivered_pilot_is_read_only_but_package_remains_readable(
    db, actors, pilot, service
):
    owner, _, operator, *_ = actors
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    pilot.status = "delivered"
    pilot.version += 1
    db.commit()
    assert service.get_manifest(owner, package.package_id).manifest
    checklist = service.get_checklist(owner, pilot.pilot_id)
    with pytest.raises(DeliveryStateConflictError):
        service.verify_item(
            operator,
            pilot.pilot_id,
            "external_report_docx",
            VersionedChecklistRequest(
                expected_checklist_version=checklist.version,
                idempotency_key="verify-delivered",
            ),
        )


def test_stale_assembling_recovery_marks_failed_and_allows_manual_retry(
    db, actors, pilot, service, clock
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    stale_time = clock.value - timedelta(minutes=11)
    assembly = DeliveryPackageAssembly(
        assembly_id="interrupted-assembly",
        pilot_id=pilot.pilot_id,
        checklist_id=checklist.checklist_id,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
        status="assembling",
        version=1,
        attempt=1,
        request_fingerprint="a" * 64,
        source_pilot_version=pilot.version,
        checklist_version=checklist.version,
        template_version=checklist.template_version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        manifest_version="delivery-manifest-v1",
        created_by_id=operator.id,
        created_at=stale_time,
        updated_at=stale_time,
        started_at=stale_time,
    )
    db.add(assembly)
    db.commit()
    assert service.recover_interrupted_assemblies(operator) == 1
    db.refresh(assembly)
    assert assembly.status == "failed"
    assert assembly.failure_code == "assembly_interrupted"
    result = service.retry_assembly(
        operator,
        assembly.assembly_id,
        RetryAssemblyRequest(
            expected_assembly_version=assembly.version,
            idempotency_key="retry-interrupted",
        ),
    )
    assert result.status == AssemblyStatus.ASSEMBLED
    assert result.attempt == 2


def test_router_contract_hides_stack_traces_and_registers_download(
    db, actors, pilot, service
):
    owner, other, operator, *_ = actors
    create_checklist(service, operator, pilot)
    app = FastAPI()
    app.include_router(router)

    def override_db():
        yield db

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_current_user] = lambda: other
    client = TestClient(app)
    response = client.get(f"/pilot-delivery/pilots/{pilot.pilot_id}/checklist")
    assert response.status_code == 403
    body = response.json()
    assert body["detail"]["code"] == "delivery_access_denied"
    assert "Traceback" not in response.text
    assert str(service.archive_root) not in response.text

    app.dependency_overrides[get_current_user] = lambda: owner
    response = client.get(f"/pilot-delivery/pilots/{pilot.pilot_id}/checklist")
    assert response.status_code == 200
    assert "pilot_ref" in response.json()
    assert pilot.pilot_id not in response.text


def test_synchronous_assembly_route_uses_fastapi_threadpool_boundary():
    assert not python_inspect.iscoroutinefunction(request_assembly_route)


def test_package_and_audit_lists_use_stable_bounded_cursors(db, actors, pilot, service):
    owner, _, operator, *_ = actors
    assemble(service, operator, pilot)
    checklist = service.get_checklist(operator, pilot.pilot_id)
    for item_key in (
        "external_report_docx",
        "one_page_summary",
        "provenance_summary",
    ):
        checklist = service.verify_item(
            operator,
            pilot.pilot_id,
            item_key,
            VersionedChecklistRequest(
                expected_checklist_version=checklist.version,
                idempotency_key=f"second-{item_key}",
            ),
        )
    readiness = service.get_readiness(operator, pilot.pilot_id)
    service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="second-assembly",
        ),
    )

    first = service.list_packages(owner, pilot.pilot_id, limit=1)
    second = service.list_packages(
        owner, pilot.pilot_id, cursor=first.next_cursor, limit=1
    )
    assert len(first.items) == len(second.items) == 1
    assert first.items[0].package_id != second.items[0].package_id
    assert first.items[0].source_pilot_version == pilot.version
    assert first.items[0].checklist_template_version == 1
    assert first.next_cursor is not None
    assert second.next_cursor is None

    audit_first = service.get_audit_history(operator, pilot.pilot_id, limit=2)
    audit_second = service.get_audit_history(
        operator, pilot.pilot_id, cursor=audit_first.next_cursor, limit=2
    )
    assert len(audit_first.items) == len(audit_second.items) == 2
    assert {item.event_id for item in audit_first.items}.isdisjoint(
        item.event_id for item in audit_second.items
    )
    with pytest.raises(DeliveryValidationError):
        service.list_packages(owner, pilot.pilot_id, cursor="not-a-cursor")
    with pytest.raises(DeliveryValidationError):
        service.get_audit_history(operator, pilot.pilot_id, limit=101)


def test_artifact_change_after_validation_fails_without_partial_package(
    monkeypatch, tmp_path, db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    artifact_path = build_project_result_archive_paths(
        archive_root=service.archive_root,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
    )["external_docx"]
    changed_path = tmp_path / "changed.docx"
    shutil.copy2(artifact_path, changed_path)
    changed = Document(changed_path)
    changed.add_paragraph("Synthetic changed content.")
    changed.save(changed_path)
    changed_bytes = changed_path.read_bytes()
    original_read_bytes = Path.read_bytes
    changed_once = False

    def racing_read_bytes(path: Path) -> bytes:
        nonlocal changed_once
        if path == artifact_path and not changed_once:
            changed_once = True
            path.write_bytes(changed_bytes)
        return original_read_bytes(path)

    monkeypatch.setattr(Path, "read_bytes", racing_read_bytes)
    result = service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="artifact-race",
        ),
    )
    assert result.status == AssemblyStatus.FAILED
    assert result.failure_code == "integrity_check_failed"
    assert db.query(DeliveryPackage).count() == 0
    assert not list(service.archive_root.rglob("*.partial"))


def test_artifact_deletion_after_validation_fails_without_package(
    monkeypatch, db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    artifact_path = build_project_result_archive_paths(
        archive_root=service.archive_root,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
    )["external_docx"]
    original_read_bytes = Path.read_bytes
    deleted_once = False

    def deleting_read_bytes(path: Path) -> bytes:
        nonlocal deleted_once
        if path == artifact_path and not deleted_once:
            deleted_once = True
            path.unlink()
        return original_read_bytes(path)

    monkeypatch.setattr(Path, "read_bytes", deleting_read_bytes)
    result = service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="artifact-deletion-race",
        ),
    )
    assert result.status == AssemblyStatus.FAILED
    assert result.failure_code == "assembly_storage_error"
    assert db.query(DeliveryPackage).count() == 0
    assert not list(service.archive_root.rglob("*.partial"))


def test_recovery_claim_is_fenced_and_does_not_recover_recent_work(
    database, db, actors, pilot, service, clock, tmp_path
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    recent_time = clock.value - timedelta(minutes=9)
    assembly = DeliveryPackageAssembly(
        assembly_id="recovery-fence",
        pilot_id=pilot.pilot_id,
        checklist_id=checklist.checklist_id,
        project_id=pilot.project_id,
        run_id=pilot.run_id,
        status="assembling",
        version=1,
        attempt=1,
        request_fingerprint="d" * 64,
        source_pilot_version=pilot.version,
        checklist_version=checklist.version,
        template_version=checklist.template_version,
        artifact_set_fingerprint=readiness.artifact_set_fingerprint,
        manifest_version="delivery-manifest-v1",
        created_by_id=operator.id,
        created_at=recent_time,
        updated_at=recent_time,
        started_at=recent_time,
    )
    db.add(assembly)
    db.commit()
    assert service.recover_interrupted_assemblies(operator) == 0
    assembly.updated_at = clock.value - timedelta(minutes=11)
    db.commit()
    _, factory = database
    barrier = threading.Barrier(2)
    recovery_now = clock.value

    def recover() -> int:
        session = factory()
        try:
            concurrent_operator = (
                session.query(User).filter(User.id == operator.id).one()
            )
            concurrent_service = DeliveryPackageService(
                session,
                archive_root=tmp_path / "project-results",
                now_provider=lambda: recovery_now,
            )
            barrier.wait(timeout=5)
            return concurrent_service.recover_interrupted_assemblies(
                concurrent_operator
            )
        finally:
            session.close()

    with ThreadPoolExecutor(max_workers=2) as pool:
        results = list(pool.map(lambda _index: recover(), range(2)))
    assert sorted(results) == [0, 1]
    db.expire_all()
    assert (
        db.query(DeliveryAuditEvent)
        .filter(
            DeliveryAuditEvent.assembly_id == assembly.assembly_id,
            DeliveryAuditEvent.event_type == "package_assembly_failed",
        )
        .count()
        == 1
    )


def test_recovery_fences_old_worker_and_quarantines_published_orphan(
    database, db, actors, pilot, service, clock, tmp_path, monkeypatch
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    _, factory = database
    recovery_db = factory()
    original_build = service._build_archive
    try:
        recovery_operator = recovery_db.query(User).filter(User.id == operator.id).one()
        recovery = DeliveryPackageService(
            recovery_db,
            archive_root=tmp_path / "project-results",
            now_provider=clock,
        )

        def build_then_recover(actor, assembly):
            built = original_build(actor, assembly)
            clock.value += timedelta(minutes=11)
            assert recovery.recover_interrupted_assemblies(recovery_operator) == 1
            return built

        monkeypatch.setattr(service, "_build_archive", build_then_recover)
        with pytest.raises(DeliveryStorageError):
            service.request_assembly(
                operator,
                pilot.pilot_id,
                RequestAssemblyRequest(
                    expected_pilot_version=pilot.version,
                    expected_checklist_version=checklist.version,
                    artifact_set_fingerprint=readiness.artifact_set_fingerprint,
                    idempotency_key="late-worker",
                ),
            )
        db.expire_all()
        stored = db.query(DeliveryPackageAssembly).one()
        assert stored.status == AssemblyStatus.FAILED.value
        assert stored.failure_code == "assembly_interrupted"
        assert db.query(DeliveryPackage).count() == 0
        package_root = (
            build_project_result_archive_paths(
                archive_root=service.archive_root,
                project_id=pilot.project_id,
                run_id=pilot.run_id,
            )["dir"]
            / "delivery_packages"
        )
        assert not list(package_root.glob("*.zip"))
        assert len(list((package_root / ".orphaned").glob("*.orphan"))) == 1
    finally:
        recovery_db.close()


def test_completion_commit_failure_rolls_back_db_and_quarantines_package(
    monkeypatch, db, actors, pilot, service
):
    operator = actors[2]
    checklist, readiness = ready_checklist(service, operator, pilot)
    original_commit = db.commit
    commit_count = 0

    def fail_completion_commit():
        nonlocal commit_count
        commit_count += 1
        if commit_count == 3:
            raise SQLAlchemyError("synthetic completion failure")
        return original_commit()

    monkeypatch.setattr(db, "commit", fail_completion_commit)
    result = service.request_assembly(
        operator,
        pilot.pilot_id,
        RequestAssemblyRequest(
            expected_pilot_version=pilot.version,
            expected_checklist_version=checklist.version,
            artifact_set_fingerprint=readiness.artifact_set_fingerprint,
            idempotency_key="completion-commit-failure",
        ),
    )
    assert result.status == AssemblyStatus.FAILED
    assert result.failure_code == "assembly_storage_error"
    assert db.query(DeliveryPackage).count() == 0
    package_root = (
        build_project_result_archive_paths(
            archive_root=service.archive_root,
            project_id=pilot.project_id,
            run_id=pilot.run_id,
        )["dir"]
        / "delivery_packages"
    )
    assert not list(package_root.glob("*.zip"))
    assert len(list((package_root / ".orphaned").glob("*.orphan"))) == 1


def test_download_reference_concurrent_consumption_allows_one_success(
    database, db, actors, pilot, service, tmp_path
):
    operator = actors[2]
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    reference = service.create_download_reference(operator, package.package_id)
    _, factory = database
    barrier = threading.Barrier(2)
    now = datetime(2026, 7, 20, 0, 5, tzinfo=timezone.utc)

    class RacingDownloadService(DeliveryPackageService):
        def _verify_package_file(self, stored_package):
            path = super()._verify_package_file(stored_package)
            barrier.wait(timeout=5)
            return path

    def redeem() -> str:
        session = factory()
        try:
            concurrent_operator = (
                session.query(User).filter(User.id == operator.id).one()
            )
            concurrent_service = RacingDownloadService(
                session,
                archive_root=tmp_path / "project-results",
                now_provider=lambda: now,
            )
            try:
                concurrent_service.redeem_download_reference(
                    concurrent_operator, reference.download_reference
                )
                return "success"
            except DeliveryStateConflictError:
                return "consumed"
        finally:
            session.close()

    with ThreadPoolExecutor(max_workers=2) as pool:
        results = list(pool.map(lambda _index: redeem(), range(2)))
    assert sorted(results) == ["consumed", "success"]


def test_download_response_uses_safe_content_disposition(
    monkeypatch, db, actors, pilot, service
):
    operator = actors[2]
    assemble(service, operator, pilot)
    package = db.query(DeliveryPackage).one()
    reference = service.create_download_reference(operator, package.package_id)
    app = FastAPI()
    app.include_router(router)

    def override_db():
        yield db

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_current_user] = lambda: operator
    monkeypatch.setattr(delivery_router_module, "_service", lambda _db: service)
    response = TestClient(app).get(
        f"/pilot-delivery/downloads/{reference.download_reference}"
    )
    assert response.status_code == 200
    disposition = response.headers["content-disposition"]
    assert "pilot-delivery-package.zip" in disposition
    assert "\r" not in disposition and "\n" not in disposition
    assert str(service.archive_root) not in disposition
