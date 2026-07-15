import os
from copy import deepcopy
from pathlib import Path

import pytest

from mellow_link.services.doc_service import (
    DocumentRequest,
    DocumentService,
    DocumentType,
)
from mellow_link.services.project_results.docx_polish import (
    append_internal_review_appendix,
    build_docx_polish_report,
)


def _sample_pkg() -> dict:
    repeated_reason = "선택 이유: 주문 상태 변경 로직과 조회 SQL이 강하게 결합되어 단계적 분리가 필요합니다."
    return {
        "project": {
            "project_name": "sample_a_legacy_order_review",
            "client_name": "si company",
            "goal": "주문, 결제, 배송 상태 흐름을 기준으로 현대화 검토",
        },
        "provenance": {
            "run_id": "run-polish-001",
            "generated_at": "2026-05-02T09:00:00Z",
            "run_status": "completed",
            "app_version": "test",
            "module_version": "test-module",
            "template_key": "pilot",
            "input_assets": [
                {"name": "sample_a_order_query.sql", "size": 1234},
                {"name": "OrderService.java", "size": 5678},
            ],
        },
        "core_conclusion": (
            "청구 조정 기능은 주문 상태 변경과 결제 확인 로직이 한 흐름에 섞여 있어 유지보수 리스크가 큽니다. "
            "현행 화면과 SQL 기준으로는 원가체계 재정의보다 상태 전이와 예외 처리 기준을 먼저 정리해야 합니다. "
            "추가 검토 전까지는 부분 개선 범위를 작게 잡아야 합니다."
        ),
        "analysis_summary": [
            "주문 조회 SQL과 상태 변경 로직이 같은 흐름에 묶여 있습니다.",
            "원가계산 기준이 아니라 주문 상태 전이 기준 확인이 필요합니다.",
        ],
        "structure_comparison": {
            "current_structure": [
                "주문 조회, 결제 확인, 배송 상태 갱신이 한 서비스 흐름에 모여 있습니다.",
            ]
        },
        "decision_items": [
            {
                "statement": "청구 조정 기능의 예외 처리가 주문 상태 보정과 섞여 있습니다."
            },
            {
                "statement": "선택 근거: 주문 상태 변경 로직과 조회 SQL이 강하게 결합되어 단계적 분리가 필요합니다."
            },
        ],
        "design_options": [
            {
                "name": "A",
                "description": "현행 구조 유지",
                "pros": ["빠르게 검토 가능"],
                "cons": ["구조 개선 효과 제한"],
            },
            {
                "name": "B",
                "description": "주문 상태 흐름을 부분 분리",
                "pros": ["리스크 통제"],
                "cons": ["범위 합의 필요"],
            },
        ],
        "recommended_option": {
            "name": "B",
            "description": "주문 상태 흐름을 기준으로 부분 개선을 우선 검토합니다.",
            "selection_reason": repeated_reason,
        },
        "risks": [
            "청구 조정 기능 범위를 그대로 두면 상태 보정 기준이 모호해집니다.",
            "제조경비 자료가 없어 판단이 제한됩니다.",
        ],
        "execution_plan": [
            {
                "phase": "1단계",
                "action": "주문 상태 전이 기준 확인",
                "tasks": ["주문 SQL 검토", "상태 변경 예외 확인"],
                "deliverables": ["검토 메모"],
            }
        ],
        "report_questions": ["주문 상태 변경 기준은 어디에서 관리되는가?"],
    }


def _summary_lines(markdown: str) -> list[str]:
    lines = markdown.splitlines()
    start = lines.index("## 1. 1페이지 요약") + 1
    end = lines.index("## 2. 분석 범위와 입력 자료")
    return [line for line in lines[start:end] if line.strip()]


def test_polish_summary_is_five_lines_or_less_and_compressed():
    markdown = build_docx_polish_report(_sample_pkg())
    summary = _summary_lines(markdown)

    assert len(summary) <= 5
    assert all(len(line) < 220 for line in summary)
    assert "원가체계" not in "\n".join(summary)


def test_polish_removes_domain_pollution_and_demo_names():
    markdown = build_docx_polish_report(_sample_pkg(), surface_mode="external")

    assert "청구 조정" not in markdown
    assert "원가계산" not in markdown
    assert "제조경비" not in markdown
    assert "sample_a" not in markdown
    assert "si company" not in markdown.lower()
    assert "주문관리 현대화 검토 사례" in markdown
    assert "고객사 A" in markdown
    assert "상태 보정" in markdown


def test_polish_deduplicates_selection_reasons():
    markdown = build_docx_polish_report(_sample_pkg())

    assert markdown.count("주문 상태 변경 로직과 조회 SQL이 강하게 결합") == 1
    assert "선택 및 권장 근거" in markdown


def test_polish_includes_report_tables():
    markdown = build_docx_polish_report(_sample_pkg())

    assert "| 유형 | 파일명 | 크기 |" in markdown
    assert "| 선택지 | 설명 | 장점 | 단점 |" in markdown
    assert "| 리스크 | 영향 | 대응 |" in markdown
    assert "| 단계 | 작업 | 비고 |" in markdown


def test_polish_uses_normal_result_package_structure_summary():
    pkg = _sample_pkg()
    pkg["design_options"] = [
        {
            "name": "A",
            "structure_summary": "주문 상태 변경 책임을 서비스 경계로 분리",
            "advantages": ["변경 영향 범위가 명확함"],
            "risks": ["전환 순서 합의 필요"],
        }
    ]

    markdown = build_docx_polish_report(pkg)

    assert "주문 상태 변경 책임을 서비스 경계로 분리" in markdown
    assert "| A | 개선 선택지 |" not in markdown


def test_polish_separates_internal_and_external_docx_surface():
    internal = build_docx_polish_report(_sample_pkg(), surface_mode="internal")
    external = build_docx_polish_report(_sample_pkg(), surface_mode="external")

    assert "run-polish-001" in internal
    assert "OrderService.java" in internal
    assert "주문관리 현대화 검토 사례_order_query.sql" in internal
    assert "run-polish-001" not in external
    assert "sample_a_order_query.sql" not in external
    assert "| 유형 | 설명 |" in external
    assert "상세 추적 정보와 입력 파일명은 내부 검토본에서 확인합니다." in external
    assert internal != external


def test_polish_output_is_deterministic_and_keeps_fixed_section_order():
    first = build_docx_polish_report(_sample_pkg())
    second = build_docx_polish_report(_sample_pkg())
    expected_sections = [
        "## 1. 1페이지 요약",
        "## 2. 분석 범위와 입력 자료",
        "## 3. 현행 구조/업무 흐름 요약",
        "## 4. 핵심 문제",
        "## 5. 개선 선택지",
        "## 6. 권장안",
        "## 7. 리스크와 검토 필요 사항",
        "## 8. 단계별 실행 준비 계획",
        "## 9. 분석 근거와 provenance",
    ]

    assert first == second
    assert [first.index(section) for section in expected_sections] == sorted(
        first.index(section) for section in expected_sections
    )


def test_polish_handles_long_summary_without_variable_width_lookbehind():
    pkg = _sample_pkg()
    pkg["core_conclusion"] = (
        "주문 조회와 상태 변경이 하나의 서비스에 결합되어 있습니다. "
        "결제 확인과 배송 갱신도 같은 트랜잭션 경계에 포함되어 있습니다. "
        "예외 처리 기준과 재처리 정책은 입력 자료에서 충분히 확인되지 않습니다. "
        "따라서 작은 범위의 구조 분리와 검증 계획을 먼저 합의해야 합니다."
    )

    markdown = build_docx_polish_report(pkg)

    assert (
        "- 한 줄 결론: 주문 조회와 상태 변경이 하나의 서비스에 결합되어 있습니다."
        in markdown
    )


def test_external_polish_excludes_internal_provenance_and_source_details():
    pkg = deepcopy(_sample_pkg())
    pkg["provenance"].update(
        {
            "safe_bundle_id": "safe-bundle-internal-001",
            "original_path": "C:/private/customer/order.sql",
            "mapping_path": "C:/private/customer/mapping.json",
        }
    )
    pkg["internal_canonical"] = {"raw_content": "private source content"}
    pkg["project"]["goal"] = "Review C:/private/customer/mapping.json"
    pkg["core_conclusion"] = "private source content must not leave run-polish-001."
    pkg["risks"].append("OrderService.java contains asset-private-002.")

    external = build_docx_polish_report(pkg, surface_mode="external")

    for internal_value in (
        "run-polish-001",
        "safe-bundle-internal-001",
        "C:/private/customer/order.sql",
        "C:/private/customer/mapping.json",
        "private source content",
        "OrderService.java",
        "asset-private-002",
    ):
        assert internal_value not in external
    assert "## 9. 산출물 기준" in external
    assert "## 9. 분석 근거와 provenance" not in external


@pytest.mark.asyncio
async def test_external_pilot_report_generates_reopenable_docx(tmp_path: Path):
    service = DocumentService(output_dir=tmp_path)
    await service.initialize()
    request = DocumentRequest(
        content=build_docx_polish_report(_sample_pkg(), surface_mode="external"),
        output_type=DocumentType.DOCX,
        title="현대화 판단 보고서",
        filename="pilot-external-review.docx",
    )

    result = await service.generate(request)

    from docx import Document

    doc = Document(str(result.output_path))
    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    expected_sections = [
        "1. 1페이지 요약",
        "2. 분석 범위와 입력 자료",
        "3. 현행 구조/업무 흐름 요약",
        "4. 핵심 문제",
        "5. 개선 선택지",
        "6. 권장안",
        "7. 리스크와 검토 필요 사항",
        "8. 단계별 실행 준비 계획",
        "9. 산출물 기준",
    ]
    assert all(section in document_text for section in expected_sections)
    assert len(doc.tables) == 4
    assert "run-polish-001" not in document_text
    assert "OrderService.java" not in document_text

    artifact_dir = os.environ.get("PILOT_DOCX_ARTIFACT_DIR")
    if artifact_dir:
        destination = Path(artifact_dir) / "pilot-external-review.docx"
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(Path(result.output_path).read_bytes())
    await service.shutdown()


def test_internal_full_docx_export_keeps_review_appendix():
    report = "# 현대화 판단 보고서\n"
    appendix = "## 참고 구조 비교\n\n- synthetic-review-sentinel\n"

    full_content = append_internal_review_appendix(
        report,
        appendix,
        surface_mode="internal",
        internal_export_mode="full",
    )
    deck_only_content = append_internal_review_appendix(
        report,
        appendix,
        surface_mode="internal",
        internal_export_mode="deck-only",
    )
    external_content = append_internal_review_appendix(
        report,
        appendix,
        surface_mode="external",
        internal_export_mode="full",
    )

    assert "## 참고 구조 비교" in full_content
    assert "synthetic-review-sentinel" in full_content
    assert deck_only_content == report
    assert external_content == report


@pytest.mark.asyncio
async def test_document_service_renders_markdown_tables_as_docx_tables(tmp_path: Path):
    service = DocumentService(output_dir=tmp_path)
    await service.initialize()
    request = DocumentRequest(
        content="# 현대화 판단 보고서\n\n| 유형 | 설명 |\n| --- | --- |\n| SQL | 주문 조회 쿼리 |",
        output_type=DocumentType.DOCX,
        title="현대화 판단 보고서",
        filename="table-test.docx",
    )

    result = await service.generate(request)

    from docx import Document

    doc = Document(str(result.output_path))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "유형"
    assert doc.tables[0].cell(1, 1).text == "주문 조회 쿼리"
    await service.shutdown()
