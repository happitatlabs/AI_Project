from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from mellow_link.services.doc_service import (
    DocumentRequest,
    DocumentService,
    DocumentType,
)
from mellow_link.services.project_results.archive import (
    build_project_result_archive_paths,
    persist_project_result_archive,
)
from mellow_link.services.project_results.docx_polish import build_docx_polish_report


class _RecordingLogger:
    def __init__(self) -> None:
        self.warnings: list[tuple[object, ...]] = []

    def warning(self, *args: object) -> None:
        self.warnings.append(args)


@pytest.mark.asyncio
async def test_project_result_archive_uses_deterministic_paths_and_reopenable_docx(
    tmp_path: Path,
):
    project = SimpleNamespace(
        id="proj_pilot_delivery",
        project_name="합성 파일럿 프로젝트",
        client_name="가상 고객사",
    )
    result_package = {
        "project": {
            "project_name": project.project_name,
            "client_name": project.client_name,
            "goal": "프로젝트 결과 보관 검증",
        },
        "provenance": {
            "run_id": "run_pilot_delivery",
            "generated_at": "2026-07-14T00:00:00Z",
            "input_assets": [{"name": "synthetic.sql", "size": 100}],
        },
        "analysis_summary": ["합성 입력의 업무 흐름을 검토합니다."],
        "risks": ["추가 자료 확인이 필요합니다."],
    }
    service = DocumentService(output_dir=tmp_path / "generated")
    await service.initialize()
    logger = _RecordingLogger()

    async def generate_docx(project, package, *, surface_mode, **kwargs):
        generated = await service.generate(
            DocumentRequest(
                content=build_docx_polish_report(package, surface_mode=surface_mode),
                output_type=DocumentType.DOCX,
                title="현대화 판단 보고서",
                filename="pilot-result.docx",
            )
        )
        return generated.output_path, "pilot-result.docx"

    result = await persist_project_result_archive(
        project,
        run_id="run_pilot_delivery",
        db=object(),
        archive_root=tmp_path / "archive",
        logger=logger,
        get_run_snapshot_fn=lambda *args, **kwargs: {"status": "completed"},
        get_run_events_fn=lambda *args, **kwargs: [{"type": "run_finished"}],
        extract_structured_result_fn=lambda events: {"status": "completed"},
        build_assets_payload_fn=lambda *args, **kwargs: [],
        build_result_package_fn=lambda *args, **kwargs: {},
        extract_polish_bundle_fn=lambda *args, **kwargs: None,
        result_package_markdown_fn=lambda *args, **kwargs: "# 현대화 판단 보고서\n",
        generate_result_package_docx_fn=generate_docx,
        result_package=result_package,
    )

    expected_dir = tmp_path / "archive" / "proj_pilot_delivery" / "run_pilot_delivery"
    expected_markdown = expected_dir / "result.md"
    expected_docx = expected_dir / "result.docx"
    expected_external_docx = expected_dir / "external_result.docx"

    assert result == {
        "markdown_path": str(expected_markdown),
        "docx_path": str(expected_docx),
        "external_docx_path": str(expected_external_docx),
    }
    assert expected_markdown.read_text(encoding="utf-8") == "# 현대화 판단 보고서\n"
    assert expected_docx.is_file()
    reopened = Document(expected_docx)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    assert "1. 1페이지 요약" in text
    assert "8. 단계별 실행 준비 계획" in text
    assert "9. 분석 근거와 provenance" in text
    assert len(reopened.tables) == 4
    assert logger.warnings == []
    external = Document(expected_external_docx)
    external_text = "\n".join(paragraph.text for paragraph in external.paragraphs)
    assert "9. 산출물 기준" in external_text
    assert "9. 분석 근거와 provenance" not in external_text
    assert "run_pilot_delivery" not in external_text

    same_paths = build_project_result_archive_paths(
        archive_root=tmp_path / "archive",
        project_id=project.id,
        run_id="run_pilot_delivery",
    )
    other_run_paths = build_project_result_archive_paths(
        archive_root=tmp_path / "archive",
        project_id=project.id,
        run_id="run_pilot_delivery_2",
    )
    assert same_paths["docx"] == expected_docx
    assert other_run_paths["docx"] != expected_docx
    await service.shutdown()
